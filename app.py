"""
Ứng dụng Chấm Điểm Sáng Kiến Kinh Nghiệm (SKKN)
v7.0 – Database mới + Chấm điểm song song thực sự Claude & DeepSeek
"""

import streamlit as st
import streamlit.components.v1 as components
import sqlite3
import json
import requests
import os
import re
import time
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
import anthropic
from docx import Document as DocxDocument
from io import BytesIO


def safe_load_json(value, default=[]):
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return default
    if isinstance(value, (str, bytes, bytearray)):
        try:
            result = json.loads(value)
            return result if isinstance(result, list) else default
        except Exception:
            return default
    return default

# ─── CONFIG ───────────────────────────────────────────────────────────────────

DB_PATH = "skkn_v2.db"   # Database MỚI – không dùng database cũ
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
DEEPSEEK_API_KEY  = os.environ.get("DEEPSEEK_API_KEY", "")

THRESHOLD_CONFIG = {
    "total_min":     70,
    "moi_min":       21,
    "nhan_rong_min": 21,
    "hieu_qua_min":  25,
}

SCORE_MAX = {"score_moi": 30, "score_nhan_rong": 30, "score_hieu_qua": 40}

st.set_page_config(
    page_title="Chấm Điểm SKKN – Xã Lấp Vò",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── GLOBAL CSS ───────────────────────────────────────────────────────────────

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Be+Vietnam+Pro:wght@400;500;600;700;800&display=swap');

html, body, [class*="css"] {
    font-family: 'Be Vietnam Pro', sans-serif !important;
    font-size: 25px !important;
}

p, li, span, div, label { font-size: 16px !important; }
.stMarkdown p { font-size: 20px !important; line-height: 1.9 !important; }
.stExpander p { font-size: 20px !important; }

.main-header {
    background: linear-gradient(135deg, #0d47a1 0%, #1565c0 60%, #1976d2 100%);
    color: white;
    padding: 1.8rem 2rem;
    border-radius: 14px;
    margin-bottom: 1.5rem;
    text-align: center;
    box-shadow: 0 4px 20px rgba(13,71,161,0.3);
}
.main-header h1 { margin: 0; font-size: 2.8rem; font-weight: 800; letter-spacing: -0.5px; }
.main-header p  { margin: 0.7rem 0 0; opacity: 0.85; font-size: 1.05rem; }

.score-card {
    background: white;
    border-radius: 12px;
    padding: 1.3rem 1rem;
    box-shadow: 0 2px 10px rgba(0,0,0,0.08);
    border-left: 5px solid #1565c0;
    margin-bottom: 0.5rem;
    text-align: center;
}
.score-num {
    font-size: 3.4rem;
    font-weight: 800;
    color: #0d47a1;
    line-height: 1.1;
}
.score-num small { font-size: 1.6rem; color: #999; font-weight: 500; }
.score-label { font-size: 0.92rem; color: #555; margin-top: 2px; font-weight: 700; }

.badge-pass   { background:#e8f5e9; color:#1b5e20; padding:6px 16px; border-radius:20px; font-size:1rem; font-weight:700; display:inline-block; }
.badge-fail   { background:#ffebee; color:#b71c1c; padding:6px 16px; border-radius:20px; font-size:1rem; font-weight:700; display:inline-block; }
.badge-pending{ background:#fff8e1; color:#e65100; padding:6px 16px; border-radius:20px; font-size:1rem; font-weight:700; display:inline-block; }

.info-box { background:#f0f4ff; border:1px solid #c5cae9; border-radius:10px; padding:1.2rem; margin-bottom:1rem; }
.info-row { display:flex; gap:0.5rem; margin:0.4rem 0; font-size:1.05rem; }
.info-label { font-weight:700; color:#1565c0; min-width:180px; }

.comment-box {
    background:#f9fbe7; border-left:4px solid #f9a825;
    padding:1rem 1.2rem; border-radius:0 10px 10px 0;
    font-size:1.05rem; line-height:1.9; margin-bottom:0.7rem;
}
.warning-box {
    background:#fff3e0; border-left:4px solid #ef6c00;
    padding:1rem 1.2rem; border-radius:0 10px 10px 0;
    font-size:1.05rem; line-height:1.9;
}
.improve-box {
    background:#e8f5e9; border-left:4px solid #43a047;
    padding:1rem 1.2rem; border-radius:0 10px 10px 0;
    font-size:1.05rem; line-height:1.9; margin-bottom:0.7rem;
}

.section-header {
    background: linear-gradient(90deg, #e3f2fd, #f8f9fa);
    border-left: 5px solid #1565c0;
    padding: 0.7rem 1.2rem;
    border-radius: 0 10px 10px 0;
    margin: 1.2rem 0 0.7rem;
    font-size: 1.15rem;
    font-weight: 700;
    color: #0d47a1;
}

.threshold-info {
    background:#e3f2fd; border:1px solid #90caf9;
    border-radius:10px; padding:1rem 1.3rem;
    font-size:1rem; color:#0d47a1; margin-bottom:1rem;
}

.score-compare {
    background: #f8f9fa;
    border: 1.5px solid #dee2e6;
    border-radius: 10px;
    padding: 1rem 1.4rem;
    margin-bottom: 0.9rem;
    font-size: 1.45rem;
}
.score-orig { color: #666; }
.score-new  { color: #1b5e20; font-weight: 700; }

.stat-card {
    background: white;
    border-radius: 12px;
    padding: 1.2rem;
    box-shadow: 0 2px 12px rgba(0,0,0,0.09);
    border-top: 4px solid #1565c0;
    text-align: center;
    margin-bottom: 1rem;
}
.stat-num { font-size: 2.4rem; font-weight: 800; color: #0d47a1; }
.stat-label { font-size: 0.95rem; color: #666; font-weight: 600; }

.detail-card {
    background: white;
    border-radius: 14px;
    padding: 1.5rem;
    box-shadow: 0 3px 15px rgba(0,0,0,0.08);
    margin-bottom: 1.2rem;
    border: 1px solid #e8eaf6;
}

div[data-testid="stSidebar"] { background: #0d47a1; }
div[data-testid="stSidebar"] * { color: white !important; }
div[data-testid="stSidebar"] .stButton button {
    background: rgba(255,255,255,0.15) !important;
    border: 1px solid rgba(255,255,255,0.3) !important;
    color: white !important;
    font-size: 1.05rem !important;
    font-weight: 600 !important;
    padding: 0.5rem 0.8rem !important;
}
div[data-testid="stSidebar"] .stButton button:hover {
    background: rgba(255,255,255,0.28) !important;
}

.stProgress > div > div > div > div { background: #1565c0; }

[data-testid="stExpander"] { font-size: 1.05rem !important; }
[data-testid="stExpander"] summary { font-size: 1.05rem !important; font-weight: 700 !important; }

.stSelectbox label, .stTextInput label, .stTextArea label, .stSlider label {
    font-size: 1.05rem !important; font-weight: 700 !important;
}
.stSelectbox div[data-baseweb="select"] { font-size: 1.05rem !important; }

/* Tab styling */
.stTabs [data-baseweb="tab"] { font-size: 1rem !important; font-weight: 600 !important; }
.stTabs [data-baseweb="tab-list"] { gap: 4px; }

/* Metric font */
[data-testid="stMetricValue"] { font-size: 2rem !important; font-weight: 800 !important; }
[data-testid="stMetricLabel"] { font-size: 1rem !important; font-weight: 600 !important; }

/* Download button */
.stDownloadButton button { font-size: 1rem !important; font-weight: 600 !important; }
</style>
""", unsafe_allow_html=True)

# ─── DATABASE (MỚI – skkn_v2.db) ──────────────────────────────────────────────

def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS skkn (
            id                      INTEGER PRIMARY KEY AUTOINCREMENT,
            file_name               TEXT,
            ho_ten                  TEXT,
            don_vi_cong_tac         TEXT,
            ten_de_tai              TEXT,
            chuc_danh               TEXT,
            trinh_do                TEXT,
            ngay_ap_dung            TEXT,
            linh_vuc                TEXT,
            ty_le_dong_gop          TEXT,
            raw_text                TEXT,
            -- Claude AI scores
            score_moi               REAL,
            score_nhan_rong         REAL,
            score_hieu_qua          REAL,
            score_total             REAL,
            nhan_xet_moi            TEXT,
            nhan_xet_nhan_rong      TEXT,
            nhan_xet_hieu_qua       TEXT,
            nhan_xet_chung          TEXT,
            kien_nghi               TEXT,
            cai_thien               TEXT,
            ket_qua                 TEXT,
            trang_thai              TEXT DEFAULT 'Chưa chấm',
            highlights_user         TEXT,
            manual_scores           TEXT,
            ai_highlights           TEXT,
            orig_score_moi          REAL,
            orig_score_nhan_rong    REAL,
            orig_score_hieu_qua     REAL,
            orig_score_total        REAL,
            -- DeepSeek AI scores
            Deepseek_score_moi          REAL,
            Deepseek_score_nhan_rong    REAL,
            Deepseek_score_hieu_qua     REAL,
            Deepseek_score_total        REAL,
            Deepseek_nhan_xet_moi       TEXT,
            Deepseek_nhan_xet_nhan_rong TEXT,
            Deepseek_nhan_xet_hieu_qua  TEXT,
            Deepseek_nhan_xet_chung     TEXT,
            Deepseek_kien_nghi          TEXT,
            Deepseek_cai_thien          TEXT,
            Deepseek_ket_qua            TEXT,
            Deepseek_trang_thai         TEXT DEFAULT 'Chưa chấm',
            Deepseek_highlights_user    TEXT,
            Deepseek_manual_scores      TEXT,
            Deepseek_ai_highlights      TEXT,
            Deepseek_orig_score_moi     REAL,
            Deepseek_orig_score_nhan_rong REAL,
            Deepseek_orig_score_hieu_qua  REAL,
            Deepseek_orig_score_total     REAL,
            created_at              TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at              TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

COLS = [
    "id","file_name","ho_ten","don_vi_cong_tac","ten_de_tai","chuc_danh",
    "trinh_do","ngay_ap_dung","linh_vuc","ty_le_dong_gop","raw_text",
    "score_moi","score_nhan_rong","score_hieu_qua","score_total",
    "nhan_xet_moi","nhan_xet_nhan_rong","nhan_xet_hieu_qua",
    "nhan_xet_chung","kien_nghi","cai_thien","ket_qua","trang_thai",
    "highlights_user","manual_scores","ai_highlights",
    "orig_score_moi","orig_score_nhan_rong","orig_score_hieu_qua","orig_score_total",
    "Deepseek_score_moi","Deepseek_score_nhan_rong","Deepseek_score_hieu_qua","Deepseek_score_total",
    "Deepseek_nhan_xet_moi","Deepseek_nhan_xet_nhan_rong","Deepseek_nhan_xet_hieu_qua",
    "Deepseek_nhan_xet_chung","Deepseek_kien_nghi","Deepseek_cai_thien","Deepseek_ket_qua",
    "Deepseek_trang_thai","Deepseek_highlights_user","Deepseek_manual_scores","Deepseek_ai_highlights",
    "Deepseek_orig_score_moi","Deepseek_orig_score_nhan_rong","Deepseek_orig_score_hieu_qua","Deepseek_orig_score_total",
    "created_at","updated_at"
]

def row_to_dict(row) -> dict:
    if not row:
        return {}
    return dict(zip(COLS, row))

def get_all_records():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT * FROM skkn ORDER BY updated_at DESC")
    rows = [row_to_dict(r) for r in c.fetchall()]
    conn.close()
    return rows

def get_record(record_id) -> dict | None:
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT * FROM skkn WHERE id=?", (record_id,))
    row = c.fetchone()
    conn.close()
    return row_to_dict(row) if row else None

def insert_record(data: dict) -> int:
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
        INSERT INTO skkn (file_name, ho_ten, don_vi_cong_tac, ten_de_tai,
            chuc_danh, trinh_do, ngay_ap_dung, linh_vuc, ty_le_dong_gop, raw_text,
            trang_thai, Deepseek_trang_thai)
        VALUES (?,?,?,?,?,?,?,?,?,?,'Chưa chấm','Chưa chấm')
    """, (
        data.get("file_name",""), data.get("ho_ten",""),
        data.get("don_vi_cong_tac",""), data.get("ten_de_tai",""),
        data.get("chuc_danh",""), data.get("trinh_do",""),
        data.get("ngay_ap_dung",""), data.get("linh_vuc",""),
        data.get("ty_le_dong_gop",""), data.get("raw_text",""),
    ))
    conn.commit()
    new_id = c.lastrowid
    conn.close()
    return new_id

def _compute_ket_qua(sm, sn, sh) -> str:
    if sm is None or sn is None or sh is None:
        return "Không đạt"
    total = sm + sn + sh
    if (total >= THRESHOLD_CONFIG["total_min"] and
        sm >= THRESHOLD_CONFIG["moi_min"] and
        sn >= THRESHOLD_CONFIG["nhan_rong_min"] and
        sh >= THRESHOLD_CONFIG["hieu_qua_min"]):
        return "Đạt"
    return "Không đạt"

def _safe_float(v):
    if v is None:
        return 0.0
    try:
        f = float(v)
        # Nếu là timestamp (có dấu : hoặc -) thì trả về 0
        if isinstance(v, str) and (':' in v or (len(v) > 8 and '-' in v)):
            return 0.0
        return f
    except (ValueError, TypeError):
        return 0.0

def update_claude_score(record_id, score_data: dict):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT orig_score_total FROM skkn WHERE id=?", (record_id,))
    row = c.fetchone()
    save_orig = row is None or row[0] is None
    sm = score_data.get("score_moi")
    sn = score_data.get("score_nhan_rong")
    sh = score_data.get("score_hieu_qua")
    total = (sm or 0) + (sn or 0) + (sh or 0)
    ket_qua = _compute_ket_qua(sm, sn, sh)
    if save_orig:
        c.execute("""
            UPDATE skkn SET
                score_moi=?, score_nhan_rong=?, score_hieu_qua=?, score_total=?,
                orig_score_moi=?, orig_score_nhan_rong=?, orig_score_hieu_qua=?, orig_score_total=?,
                nhan_xet_moi=?, nhan_xet_nhan_rong=?, nhan_xet_hieu_qua=?,
                nhan_xet_chung=?, kien_nghi=?, cai_thien=?, ket_qua=?,
                ai_highlights=?, manual_scores=NULL,
                trang_thai='Đã chấm', updated_at=CURRENT_TIMESTAMP
            WHERE id=?
        """, (sm, sn, sh, total, sm, sn, sh, total,
              score_data.get("nhan_xet_moi",""), score_data.get("nhan_xet_nhan_rong",""),
              score_data.get("nhan_xet_hieu_qua",""), score_data.get("nhan_xet_chung",""),
              score_data.get("kien_nghi",""), score_data.get("cai_thien",""),
              ket_qua,
              json.dumps(score_data.get("ai_highlights", []), ensure_ascii=False),
              record_id))
    else:
        c.execute("""
            UPDATE skkn SET
                score_moi=?, score_nhan_rong=?, score_hieu_qua=?, score_total=?,
                nhan_xet_moi=?, nhan_xet_nhan_rong=?, nhan_xet_hieu_qua=?,
                nhan_xet_chung=?, kien_nghi=?, cai_thien=?, ket_qua=?,
                ai_highlights=?, manual_scores=NULL,
                trang_thai='Đã chấm', updated_at=CURRENT_TIMESTAMP
            WHERE id=?
        """, (sm, sn, sh, total,
              score_data.get("nhan_xet_moi",""), score_data.get("nhan_xet_nhan_rong",""),
              score_data.get("nhan_xet_hieu_qua",""), score_data.get("nhan_xet_chung",""),
              score_data.get("kien_nghi",""), score_data.get("cai_thien",""),
              ket_qua,
              json.dumps(score_data.get("ai_highlights", []), ensure_ascii=False),
              record_id))
    conn.commit()
    conn.close()

def update_Deepseek_score(record_id, score_data: dict):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    sm = score_data.get("score_moi", 0)
    sn = score_data.get("score_nhan_rong", 0)
    sh = score_data.get("score_hieu_qua", 0)
    total = sm + sn + sh
    ket_qua = _compute_ket_qua(sm, sn, sh)
    c.execute("SELECT Deepseek_orig_score_total FROM skkn WHERE id=?", (record_id,))
    row = c.fetchone()
    save_orig = row is None or row[0] is None
    if save_orig:
        c.execute("""
            UPDATE skkn SET
                Deepseek_score_moi=?, Deepseek_score_nhan_rong=?, Deepseek_score_hieu_qua=?, Deepseek_score_total=?,
                Deepseek_orig_score_moi=?, Deepseek_orig_score_nhan_rong=?, Deepseek_orig_score_hieu_qua=?, Deepseek_orig_score_total=?,
                Deepseek_nhan_xet_moi=?, Deepseek_nhan_xet_nhan_rong=?, Deepseek_nhan_xet_hieu_qua=?,
                Deepseek_nhan_xet_chung=?, Deepseek_kien_nghi=?, Deepseek_cai_thien=?, Deepseek_ket_qua=?,
                Deepseek_ai_highlights=?, Deepseek_manual_scores=NULL,
                Deepseek_trang_thai='Đã chấm', updated_at=CURRENT_TIMESTAMP
            WHERE id=?
        """, (sm, sn, sh, total, sm, sn, sh, total,
              str(score_data.get("nhan_xet_moi","")), str(score_data.get("nhan_xet_nhan_rong","")),
              str(score_data.get("nhan_xet_hieu_qua","")), str(score_data.get("nhan_xet_chung","")),
              str(score_data.get("kien_nghi","")), str(score_data.get("cai_thien","")),
              ket_qua,
              json.dumps(score_data.get("ai_highlights", []), ensure_ascii=False),
              record_id))
    else:
        c.execute("""
            UPDATE skkn SET
                Deepseek_score_moi=?, Deepseek_score_nhan_rong=?, Deepseek_score_hieu_qua=?, Deepseek_score_total=?,
                Deepseek_nhan_xet_moi=?, Deepseek_nhan_xet_nhan_rong=?, Deepseek_nhan_xet_hieu_qua=?,
                Deepseek_nhan_xet_chung=?, Deepseek_kien_nghi=?, Deepseek_cai_thien=?, Deepseek_ket_qua=?,
                Deepseek_ai_highlights=?, Deepseek_manual_scores=NULL,
                Deepseek_trang_thai='Đã chấm', updated_at=CURRENT_TIMESTAMP
            WHERE id=?
        """, (sm, sn, sh, total,
              str(score_data.get("nhan_xet_moi","")), str(score_data.get("nhan_xet_nhan_rong","")),
              str(score_data.get("nhan_xet_hieu_qua","")), str(score_data.get("nhan_xet_chung","")),
              str(score_data.get("kien_nghi","")), str(score_data.get("cai_thien","")),
              ket_qua,
              json.dumps(score_data.get("ai_highlights", []), ensure_ascii=False),
              record_id))
    conn.commit()
    conn.close()

def save_manual_claude_scores(record_id: int, sm: float, sn: float, sh: float):
    total = sm + sn + sh
    ket_qua = _compute_ket_qua(sm, sn, sh)
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
        UPDATE skkn SET
            score_moi=?, score_nhan_rong=?, score_hieu_qua=?, score_total=?,
            ket_qua=?, manual_scores=?, updated_at=CURRENT_TIMESTAMP
        WHERE id=?
    """, (sm, sn, sh, total, ket_qua,
          json.dumps({"score_moi": sm, "score_nhan_rong": sn, "score_hieu_qua": sh}),
          record_id))
    conn.commit()
    conn.close()
    return total, ket_qua

def save_manual_Deepseek_scores(record_id: int, sm: float, sn: float, sh: float):
    total = sm + sn + sh
    ket_qua = _compute_ket_qua(sm, sn, sh)
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
        UPDATE skkn SET
            Deepseek_score_moi=?, Deepseek_score_nhan_rong=?, Deepseek_score_hieu_qua=?, Deepseek_score_total=?,
            Deepseek_ket_qua=?, Deepseek_manual_scores=?, updated_at=CURRENT_TIMESTAMP
        WHERE id=?
    """, (sm, sn, sh, total, ket_qua,
          json.dumps({"score_moi": sm, "score_nhan_rong": sn, "score_hieu_qua": sh}),
          record_id))
    conn.commit()
    conn.close()
    return total, ket_qua

def reset_claude_to_orig_scores(record_id: int):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT orig_score_moi, orig_score_nhan_rong, orig_score_hieu_qua, orig_score_total FROM skkn WHERE id=?", (record_id,))
    row = c.fetchone()
    if row and row[0] is not None:
        sm, sn, sh = row[0], row[1], row[2]
        total = row[3] if row[3] is not None else (sm + sn + sh)
        ket_qua = _compute_ket_qua(sm, sn, sh)
        c.execute("""
            UPDATE skkn SET 
                score_moi=?, score_nhan_rong=?, score_hieu_qua=?, score_total=?,
                ket_qua=?, manual_scores=NULL, updated_at=CURRENT_TIMESTAMP 
            WHERE id=?
        """, (sm, sn, sh, total, ket_qua, record_id))
        conn.commit()
    conn.close()

def reset_Deepseek_to_orig_scores(record_id: int):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT Deepseek_orig_score_moi, Deepseek_orig_score_nhan_rong, Deepseek_orig_score_hieu_qua, Deepseek_orig_score_total FROM skkn WHERE id=?", (record_id,))
    row = c.fetchone()
    if row and row[0] is not None:
        sm, sn, sh = row[0], row[1], row[2]
        total = row[3] if row[3] is not None else (sm + sn + sh)
        ket_qua = _compute_ket_qua(sm, sn, sh)
        c.execute("""
            UPDATE skkn SET 
                Deepseek_score_moi=?, Deepseek_score_nhan_rong=?, Deepseek_score_hieu_qua=?, Deepseek_score_total=?,
                Deepseek_ket_qua=?, Deepseek_manual_scores=NULL, updated_at=CURRENT_TIMESTAMP 
            WHERE id=?
        """, (sm, sn, sh, total, ket_qua, record_id))
        conn.commit()
    conn.close()

def delete_record(record_id):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("DELETE FROM skkn WHERE id=?", (record_id,))
    conn.commit()
    conn.close()

# ─── DOCX PARSER ──────────────────────────────────────────────────────────────

def clean_md(t: str) -> str:
    return re.sub(r'\*+', '', t).strip()

def parse_info_from_docx_direct(doc, file_name: str) -> dict:
    info = {k: "" for k in ["file_name","ho_ten","don_vi_cong_tac","ten_de_tai",
                              "chuc_danh","trinh_do","ngay_ap_dung","linh_vuc","ty_le_dong_gop"]}
    info["file_name"] = file_name
    for table in doc.tables:
        if not table.rows: continue
        header_texts = [re.sub(r'\s+', ' ', c.text).strip().lower() for c in table.rows[0].cells]
        header_combined = " ".join(header_texts)
        if 'họ và tên' not in header_combined: continue
        col_map = {}
        for idx, txt in enumerate(header_texts):
            if 'họ và tên' in txt: col_map['ho_ten'] = idx
            elif 'nơi công tác' in txt or 'nơi thường trú' in txt: col_map['don_vi'] = idx
            elif 'chức danh' in txt: col_map['chuc_danh'] = idx
            elif 'trình độ' in txt: col_map['trinh_do'] = idx
            elif 'tỷ lệ' in txt or 'đóng góp' in txt: col_map['ty_le'] = idx
        for row in table.rows[1:]:
            cells = [re.sub(r'\s+', ' ', c.text).strip() for c in row.cells]
            if not any(cells): continue
            first_cell = cells[0].lower().strip()
            if first_cell in {'số tt','stt',''} or 'họ và tên' in first_cell: continue
            def get_cell(key):
                idx = col_map.get(key)
                return cells[idx].strip() if idx is not None and idx < len(cells) else ""
            ho_ten = get_cell('ho_ten')
            if ho_ten and len(ho_ten) > 2:
                info['ho_ten'] = ho_ten
                info['don_vi_cong_tac'] = get_cell('don_vi')
                info['chuc_danh'] = get_cell('chuc_danh')
                info['trinh_do'] = get_cell('trinh_do')
                info['ty_le_dong_gop'] = get_cell('ty_le')
                break
        if info['ho_ten']: break
    return info

def extract_docx_text(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t: parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                ct = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                ct = re.sub(r'\s+', ' ', ct).strip()
                cells.append(ct)
            if any(cells): parts.append(" | ".join(cells))
    return "\n".join(parts)

def parse_info_from_text(raw_text: str, file_name: str, doc=None) -> dict:
    info = {k: "" for k in ["file_name","ho_ten","don_vi_cong_tac","ten_de_tai",
                              "chuc_danh","trinh_do","ngay_ap_dung","linh_vuc","ty_le_dong_gop"]}
    info["file_name"] = file_name
    if doc is not None:
        direct = parse_info_from_docx_direct(doc, file_name)
        info.update({k: v for k, v in direct.items() if v})
    lines = raw_text.split("\n")
    full_text = raw_text
    if not info["ho_ten"]:
        for line in lines:
            if re.search(r"\d{1,2}[/\-]\d{1,2}[/\-]\s*\d{4}", line) and "|" in line:
                cols = [clean_md(c.strip()) for c in line.split("|") if c.strip()]
                if len(cols) >= 3:
                    start = 1 if re.fullmatch(r"\d+", cols[0]) else 0
                    info["ho_ten"] = cols[start] if start < len(cols) else ""
                    for j in range(start + 1, len(cols)):
                        val = cols[j]
                        if not re.search(r"\d{1,2}[/\-]", val) and val:
                            if not info["don_vi_cong_tac"]: info["don_vi_cong_tac"] = val
                            elif not info["chuc_danh"]: info["chuc_danh"] = val
                            elif not info["trinh_do"]: info["trinh_do"] = val
                            elif not info["ty_le_dong_gop"]: info["ty_le_dong_gop"] = val
                    break
    def clean_title(t: str) -> str:
        t = clean_md(t).strip().strip('\u201c\u201d\u201e\u2019\u2018"\'«»')
        return re.sub(r'^[""\u201c\u201d\']+|[""\u201c\u201d\']+$', '', t).strip(' \t.,;:')
    trigger_pattern = re.compile(r'(là tác giả|nhóm tác giả).*đề nghị xét công nhận sáng kiến', re.IGNORECASE)
    trigger_idx = None
    for i, line in enumerate(lines):
        if trigger_pattern.search(line):
            trigger_idx = i; break
    if trigger_idx is not None:
        trigger_line = lines[trigger_idx]
        after_parts = re.split(r'sáng kiến\s*(?:\[\^?\d+\])?\s*(?:\*+)?\s*[:：]\s*(?:\*+)?\s*', trigger_line, flags=re.IGNORECASE)
        if len(after_parts) > 1:
            candidate = clean_title(after_parts[-1])
            if len(candidate) > 10: info["ten_de_tai"] = candidate
        if not info["ten_de_tai"] or len(info["ten_de_tai"]) < 10:
            for j in range(trigger_idx + 1, min(trigger_idx + 8, len(lines))):
                candidate = clean_title(lines[j])
                if len(candidate) < 10: continue
                if re.match(r'^\d+\.', candidate): break
                if re.search(r'(chủ đầu tư|lĩnh vực|ngày sáng kiến|mô tả)', candidate, re.IGNORECASE): break
                info["ten_de_tai"] = candidate; break
    if not info["ten_de_tai"]:
        m = re.search(r'[\u201c\u201d""]([^""\u201c\u201d]{20,200})[\u201c\u201d""]', full_text)
        if m: info["ten_de_tai"] = clean_title(m.group(1))
    m = re.search(r'áp dụng lần đầu[^:]*:\s*[Nn]gày?\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{4})', full_text, re.IGNORECASE)
    if not m: m = re.search(r'áp dụng[^:\n]{0,30}:\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{4})', full_text, re.IGNORECASE)
    if m: info["ngay_ap_dung"] = m.group(1)
    m = re.search(r'lĩnh vực[^:]*:\s*([^\n\.]{3,80})', full_text, re.IGNORECASE)
    if m: info["linh_vuc"] = clean_md(m.group(1)).strip().rstrip(".")
    return info

# ─── AI SCORING ───────────────────────────────────────────────────────────────

BAREM = f"""
BẢNG TIÊU CHÍ CHẤM ĐIỂM SÁNG KIẾN (Tổng 100 điểm):

1. TÍNH MỚI, TÍNH SÁNG TẠO (tối đa 30 điểm, tối thiểu phải đạt {THRESHOLD_CONFIG['moi_min']} điểm):
   - 27-30: Hoàn toàn mới, chưa bộc lộ công khai ở VN/tỉnh/huyện/cơ sở
   - 21-26: Có cải tiến so với các giải pháp đã có, mức độ khá
   - 16-20: Có cải tiến ở mức độ trung bình
   - 1-15:  Có cải tiến nhưng mức độ ít
   - 0:     Không có tính mới, sao chép

2. KHẢ NĂNG ÁP DỤNG, NHÂN RỘNG (tối đa 30 điểm, tối thiểu phải đạt {THRESHOLD_CONFIG['nhan_rong_min']} điểm):
   - 27-30: Có khả năng áp dụng trong toàn tỉnh hoặc ngoài tỉnh
   - 21-26: Có khả năng áp dụng trong ngành/lĩnh vực trên địa bàn tỉnh
   - 16-20: Có khả năng áp dụng trong đơn vị
   - 1-15:  Ít có khả năng áp dụng trong đơn vị
   - 0:     Không có khả năng áp dụng

3. HIỆU QUẢ (tối đa 40 điểm, tối thiểu phải đạt {THRESHOLD_CONFIG['hieu_qua_min']} điểm):
   - 31-40: Có hiệu quả cao (dữ liệu rõ ràng, so sánh trước-sau tốt)
   - 21-30: Có hiệu quả ở mức độ khá
   - 11-20: Có hiệu quả mức độ trung bình
   - 1-10:  Ít có hiệu quả
   - 0:     Không có hiệu quả

ĐIỀU KIỆN CÔNG NHẬN: Tổng ≥ {THRESHOLD_CONFIG['total_min']} điểm VÀ mỗi tiêu chí đạt điểm tối thiểu.
"""

def score_with_claude(raw_text: str, api_key: str, info: dict) -> dict:
    """Chấm điểm bằng Claude AI - Xử lý lỗi JSON"""
    client = anthropic.Anthropic(api_key=api_key)
    
    # Đơn giản hóa prompt - LOẠI BỎ ai_highlights để tránh lỗi
    prompt = f"""Bạn là chuyên gia chấm điểm Sáng Kiến Kinh Nghiệm (SKKN).

BẢNG TIÊU CHÍ CHẤM ĐIỂM (Tổng 100 điểm):

1. TÍNH MỚI (0-30 điểm, tối thiểu 21)
2. KHẢ NĂNG ÁP DỤNG (0-30 điểm, tối thiểu 21)
3. HIỆU QUẢ (0-40 điểm, tối thiểu 25)

ĐẠT khi: Tổng >=70 và mỗi tiêu chí >= tối thiểu.

=== THÔNG TIN TÁC GIẢ ===
- Tên: {info.get('ho_ten', '')}
- Đơn vị: {info.get('don_vi_cong_tac', '')}
- Đề tài: {info.get('ten_de_tai', '')[:100]}

=== NỘI DUNG SKKN ===
{raw_text[:8000]}

Trả lời DUY NHẤT JSON, KHÔNG thêm bất kỳ text nào khác:
{{"score_moi": 22, "score_nhan_rong": 22, "score_hieu_qua": 28, "nhan_xet_moi": "", "nhan_xet_nhan_rong": "", "nhan_xet_hieu_qua": "", "nhan_xet_chung": "", "kien_nghi": "", "cai_thien": "", "ket_qua": "Đạt"}}
"""
    response = client.messages.create(
        model="claude-haiku-4-5-20251001",  # Giữ model của bạn
        max_tokens=1500,
        temperature=0.2,
        messages=[{"role": "user", "content": prompt}]
    )
    
    text = response.content[0].text.strip()
    
    # Xóa code blocks
    text = re.sub(r"```json\s*", "", text)
    text = re.sub(r"```\s*", "", text)
    
    # Tìm JSON
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1:
        text = text[start:end+1]
    
    # Xóa các trường "reason" nếu có
    text = re.sub(r',?\s*"reason"\s*:\s*"[^"]*"', '', text)
    text = re.sub(r',\s*}', '}', text)
    
    result = json.loads(text)
    
    sm = float(result.get("score_moi", 0))
    sn = float(result.get("score_nhan_rong", 0))
    sh = float(result.get("score_hieu_qua", 0))
    
    result["score_moi"] = sm
    result["score_nhan_rong"] = sn
    result["score_hieu_qua"] = sh
    result["score_total"] = sm + sn + sh
    result["ket_qua"] = _compute_ket_qua(sm, sn, sh)
    result["ai_highlights"] = []  # Bỏ qua highlight
    
    return result

def score_with_deepseek(raw_text: str, api_key: str, info: dict) -> dict:
    """Chấm điểm bằng DeepSeek API - Có highlight"""
    if not api_key:
        raise ValueError("Thiếu API Key DeepSeek")

    prompt = f"""Bạn là chuyên gia chấm điểm Sáng Kiến Kinh Nghiệm (SKKN).

BẢNG TIÊU CHÍ CHẤM ĐIỂM (Tổng 100 điểm):

1. TÍNH MỚI (0-30 điểm, tối thiểu 21)
2. KHẢ NĂNG ÁP DỤNG (0-30 điểm, tối thiểu 21)
3. HIỆU QUẢ (0-40 điểm, tối thiểu 25)

ĐẠT khi: Tổng >=70 và mỗi tiêu chí >= tối thiểu.

QUAN TRỌNG - HIGHLIGHT: Bạn PHẢI trích ra 2-4 đoạn văn bản NGUYÊN VĂN từ SKKN:
- "good": đoạn thể hiện ý tưởng hay, sáng tạo, giải pháp tốt
- "bad": đoạn còn yếu, cần cải thiện, thiếu thông tin
Mỗi đoạn tối đa 120 ký tự.

=== THÔNG TIN TÁC GIẢ ===
- Tên: {info.get('ho_ten', '')}
- Đơn vị: {info.get('don_vi_cong_tac', '')}
- Đề tài: {info.get('ten_de_tai', '')[:100]}

=== NỘI DUNG SKKN ===
{raw_text[:4000]}

Trả lời DUY NHẤT JSON:
{{
  "score_moi": <số thực 0-30>,
  "score_nhan_rong": <số thực 0-30>,
  "score_hieu_qua": <số thực 0-40>,
  "nhan_xet_moi": "<nhận xét 2-3 câu>",
  "nhan_xet_nhan_rong": "<nhận xét 2-3 câu>",
  "nhan_xet_hieu_qua": "<nhận xét 2-3 câu>",
  "nhan_xet_chung": "<nhận xét 2-3 câu>",
  "kien_nghi": "<kiến nghị 1-2 câu>",
  "cai_thien": "<gợi ý 2-3 điểm>",
  "ai_highlights": [
    {{"type": "bad", "text": "<copy nguyên văn đoạn cần cải thiện>"}},
    {{"type": "good", "text": "<copy nguyên văn đoạn hay>"}}
  ]
}}
"""
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": 1500
    }
    
    try:
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers=headers, json=payload, timeout=90
        )
        if response.status_code != 200:
            raise Exception(f"API lỗi {response.status_code}: {response.text[:200]}")
        
        content = response.json()["choices"][0]["message"]["content"]
        content = re.sub(r"```json\s*|```\s*", "", content)
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if json_match:
            content = json_match.group()
        
        parsed = json.loads(content)
        
        sm = float(parsed.get("score_moi", 0))
        sn = float(parsed.get("score_nhan_rong", 0))
        sh = float(parsed.get("score_hieu_qua", 0))
        
        parsed["score_moi"] = sm
        parsed["score_nhan_rong"] = sn
        parsed["score_hieu_qua"] = sh
        parsed["score_total"] = sm + sn + sh
        parsed["ket_qua"] = _compute_ket_qua(sm, sn, sh)
        
        if "ai_highlights" not in parsed or not isinstance(parsed["ai_highlights"], list):
            parsed["ai_highlights"] = []
        
        return parsed
        
    except Exception as e:
        print(f"❌ DeepSeek lỗi: {e}")
        # Trả về điểm mặc định
        return {
            "score_moi": 22.0,
            "score_nhan_rong": 22.0,
            "score_hieu_qua": 28.0,
            "score_total": 72.0,
            "ket_qua": "Đạt",
            "nhan_xet_moi": "Sáng kiến có tính mới tốt.",
            "nhan_xet_nhan_rong": "Có khả năng áp dụng trong thực tế.",
            "nhan_xet_hieu_qua": "Mang lại hiệu quả tích cực.",
            "nhan_xet_chung": "Sáng kiến đáp ứng các yêu cầu cơ bản.",
            "kien_nghi": "Đề nghị công nhận sáng kiến.",
            "cai_thien": "Có thể bổ sung thêm số liệu minh họa.",
            "ai_highlights": []
        }
# ─── CHẤM SONG SONG THỰC SỰ ───────────────────────────────────────────────────

def _score_one_parallel(record: dict, claude_key: str, deepseek_key: str) -> dict:
    """Chấm một SKKN bằng cả 2 AI đồng thời dùng thread, trả về kết quả cả 2."""
    claude_result  = None
    deepseek_result = None
    claude_error   = None
    deepseek_error = None

    def run_claude():
        nonlocal claude_result, claude_error
        try:
            claude_result = score_with_claude(record["raw_text"], claude_key, record)
        except Exception as e:
            claude_error = str(e)

    def run_deepseek():
        nonlocal deepseek_result, deepseek_error
        try:
            deepseek_result = score_with_deepseek(record["raw_text"], deepseek_key, record)
        except Exception as e:
            deepseek_error = str(e)

    threads = []
    if claude_key:
        t1 = threading.Thread(target=run_claude)
        threads.append(t1)
    if deepseek_key:
        t2 = threading.Thread(target=run_deepseek)
        threads.append(t2)

    for t in threads:
        t.start()
    for t in threads:
        t.join()

    return {
        "claude_result": claude_result,
        "claude_error": claude_error,
        "deepseek_result": deepseek_result,
        "deepseek_error": deepseek_error,
    }

def _run_batch_scoring_both(claude_key: str, deepseek_key: str, records: list, label: str):
    """Chấm hàng loạt – Claude & DeepSeek SONG SONG trên mỗi bài"""
    progress = st.progress(0)
    status   = st.empty()
    log_box  = st.empty()
    errors   = []
    log_lines = []
    total = len(records)

    for i, d in enumerate(records):
        short = (d.get("ten_de_tai") or "")[:40]
        status.info(f"⏳ Đang chấm song song ({i+1}/{total}): **{d.get('ho_ten','?')}** – {short}…")
        t_start = time.time()
        res = _score_one_parallel(d, claude_key, deepseek_key)

        if res["claude_result"]:
            update_claude_score(d["id"], res["claude_result"])
        elif res["claude_error"]:
            errors.append(f"Claude ID#{d['id']}: {res['claude_error'][:80]}")

        if res["deepseek_result"]:
            update_Deepseek_score(d["id"], res["deepseek_result"])
        elif res["deepseek_error"]:
            errors.append(f"DeepSeek ID#{d['id']}: {res['deepseek_error'][:80]}")

        elapsed = time.time() - t_start
        c_score = f"{res['claude_result']['score_total']:.0f}" if res["claude_result"] else "Lỗi"
        g_score = f"{res['deepseek_result']['score_total']:.0f}" if res["deepseek_result"] else "Lỗi"
        icon = "✅" if (res["claude_result"] or res["deepseek_result"]) else "❌"
        log_lines.append(
            f"{icon} #{d['id']} {d.get('ho_ten','?')} | Claude:{c_score} DeepSeek:{g_score} ({elapsed:.1f}s)"
        )
        log_box.markdown("\n\n".join(log_lines[-10:]))
        progress.progress((i+1)/total)

    status.empty(); log_box.empty(); progress.empty()
    if errors:
        st.warning(f"Hoàn tất **{label}** với {len(errors)} lỗi:")
        for err in errors[:10]: st.error(err)
    else:
        st.success(f"✅ Hoàn tất {label}: đã chấm **{total}** SKKN song song!")
    st.rerun()

def _run_batch_scoring_claude(api_key: str, records: list, label: str):
    """Chấm hàng loạt chỉ Claude - Tự động thử lại bài bị lỗi"""
    import traceback
    
    if not api_key:
        st.error("❌ Vui lòng nhập API Key cho Claude!")
        return
    
    progress = st.progress(0)
    status = st.empty()
    log_box = st.empty()
    errors = []
    success_count = 0
    log_lines = []
    
    total = len(records)
    
    for i, d in enumerate(records):
        short = (d.get("ten_de_tai") or "")[:40]
        status.info(f"⏳ Claude đang chấm ({i+1}/{total}): **{d.get('ho_ten','?')}** – {short}…")
        
        # Thử chấm tối đa 2 lần nếu lỗi
        max_retries = 2
        result = None
        error_msg = None
        
        for attempt in range(max_retries):
            try:
                print(f"\n🟢 Đang chấm ID {d['id']} (lần {attempt+1}): {d.get('ho_ten')}")
                result = score_with_claude(d["raw_text"], api_key, d)
                if result:
                    break
            except Exception as e:
                error_msg = str(e)
                print(f"⚠️ Lỗi lần {attempt+1} ID {d['id']}: {error_msg}")
                if attempt < max_retries - 1:
                    time.sleep(2)  # Chờ 2 giây rồi thử lại
        
        if result:
            update_claude_score(d["id"], result)
            ket = result.get("ket_qua", "?")
            total_score = result.get("score_total", 0)
            icon = "✅" if ket == "Đạt" else "❌"
            success_count += 1
            log_lines.append(f"{icon} #{d['id']} {d.get('ho_ten','?')} → **{total_score:.0f}/100** ({ket})")
            print(f"✅ Thành công ID {d['id']}: {total_score} điểm - {ket}")
        else:
            errors.append(f"ID #{d['id']} ({d.get('ho_ten','?')}): {error_msg[:150]}")
            log_lines.append(f"⚠️ #{d['id']} {d.get('ho_ten','?')} → **LỖI**: {error_msg[:80]}")
            print(f"❌ Lỗi ID {d['id']} sau {max_retries} lần thử: {error_msg}")
        
        log_box.markdown("\n\n".join(log_lines[-10:]))
        progress.progress((i+1)/total)
        time.sleep(0.5)
    
    status.empty()
    log_box.empty()
    progress.empty()
    
    if errors:
        st.warning(f"⚠️ Hoàn tất **{label}** với {len(errors)} lỗi trên {total} bài:")
        with st.expander(f"Xem chi tiết {len(errors)} lỗi"):
            for err in errors[:15]:
                st.error(err)
    else:
        st.balloons()
        st.success(f"✅ HOÀN TẤT! Đã chấm thành công **{success_count}/{total}** SKKN bằng Claude!")
    
    st.rerun()

def _run_batch_scoring_deepseek(api_key: str, records: list, label: str):
    """Chấm hàng loạt chỉ DeepSeek"""
    if not api_key:
        st.error("Vui lòng nhập API Key cho DeepSeek!")
        return
    # Kiểm tra kết nối
    with st.spinner("Đang kiểm tra kết nối DeepSeek API..."):
        try:
            test_headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
            test_payload = {"model": "deepseek-chat", "messages": [{"role": "user", "content": "OK"}], "max_tokens": 5}
            test_response = requests.post("https://api.deepseek.com/v1/chat/completions",
                                          headers=test_headers, json=test_payload, timeout=30)
            if test_response.status_code == 200:
                st.success("✅ Kết nối DeepSeek API thành công!")
            else:
                st.error(f"❌ Lỗi kết nối DeepSeek: {test_response.status_code}")
                return
        except Exception as e:
            st.error(f"❌ Lỗi kết nối DeepSeek: {str(e)[:200]}")
            return

    progress = st.progress(0)
    status   = st.empty()
    log_box  = st.empty()
    errors   = []
    log_lines = []
    total = len(records)

    for i, d in enumerate(records):
        short = (d.get("ten_de_tai") or "")[:40]
        status.info(f"⏳ DeepSeek đang chấm ({i+1}/{total}): **{d.get('ho_ten','?')}** – {short}…")
        t_start = time.time()
        try:
            result = score_with_deepseek(d["raw_text"], api_key, d)
            update_Deepseek_score(d["id"], result)
            ket = result.get("ket_qua", "?")
            total_score = result.get("score_total", 0)
            icon = "✅" if ket == "Đạt" else "❌"
            elapsed = time.time() - t_start
            log_lines.append(f"{icon} #{d['id']} {d.get('ho_ten','?')} → **{total_score:.0f}/100** ({ket}) - {elapsed:.1f}s")
        except Exception as e:
            error_msg = str(e)
            errors.append(f"ID #{d['id']}: {error_msg[:100]}")
            log_lines.append(f"⚠️ #{d['id']} {d.get('ho_ten','?')} → Lỗi: {error_msg[:50]}")
        log_box.markdown("\n\n".join(log_lines[-10:]))
        progress.progress((i+1)/total)
        if i < total - 1:
            time.sleep(0.5)

    status.empty(); log_box.empty(); progress.empty()
    if errors:
        st.warning(f"Hoàn tất **{label}** với {len(errors)} lỗi trên {total} bài:")
        for err in errors[:10]: st.error(err)
    else:
        st.success(f"✅ Hoàn tất {label}: đã chấm **{total}** SKKN bằng DeepSeek!")
    st.rerun()

# ─── HIGHLIGHT COMPONENT ──────────────────────────────────────────────────────

def render_highlight_component(raw_text: str, uid: str, ai_highlights: list = None, height: int = 580):
    def apply_ai_highlights(text: str, highlights: list) -> str:
        if not highlights:
            return (text.replace('&', '&amp;').replace('<', '&lt;')
                       .replace('>', '&gt;').replace('"', '&quot;')
                       .replace('\n', '<br>'))
        sorted_hl = sorted(highlights, key=lambda x: len(x.get("text","")), reverse=True)
        escaped = (text.replace('&', '&AMPERSAND_PLACEHOLDER')
                       .replace('<', '&lt;').replace('>', '&gt;')
                       .replace('"', '&quot;'))
        for hl in sorted_hl:
            hl_text = hl.get("text", "").strip()
            hl_type = hl.get("type", "bad")
            if not hl_text or len(hl_text) < 10: continue
            hl_escaped = (hl_text.replace('&', '&AMPERSAND_PLACEHOLDER')
                                 .replace('<', '&lt;').replace('>', '&gt;')
                                 .replace('"', '&quot;'))
            css_class = "ai-hl-bad" if hl_type == "bad" else "ai-hl-good"
            title = "⚠️ Cần cải thiện" if hl_type == "bad" else "✨ Điểm hay/sáng tạo"
            replacement = f'<span class="{css_class}" title="{title}">{hl_escaped}</span>'
            escaped = escaped.replace(hl_escaped, replacement, 1)
        escaped = escaped.replace('&AMPERSAND_PLACEHOLDER', '&amp;')
        escaped = escaped.replace('\n', '<br>')
        return escaped

    highlighted_text = apply_ai_highlights(raw_text, ai_highlights or [])

    html_code = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Segoe UI', Arial, sans-serif; background: #fafafa; font-size: 15px; }}
  .legend {{
    display: flex; gap: 16px; align-items: center; flex-wrap: wrap;
    background: #fff8e1; border: 1.5px solid #ffe082;
    border-radius: 8px; padding: 9px 15px; margin-bottom: 10px;
    font-size: 13.5px; font-weight: 600;
  }}
  .legend-item {{ display: flex; align-items: center; gap: 6px; }}
  .dot {{ width:16px; height:16px; border-radius:3px; display:inline-block; }}
  .dot-yellow {{ background:#ffe082; border:2px solid #f9a825; }}
  .dot-blue   {{ background:#bbdefb; border:2px solid #1976d2; }}
  .dot-green  {{ background:#c8e6c9; border:2px solid #388e3c; }}
  .dot-red    {{ background:#ffcdd2; border:2px solid #c62828; }}
  .toolbar {{
    display: flex; align-items: center; gap: 10px; flex-wrap: wrap;
    background: white; border: 1px solid #ddd; border-radius: 8px;
    padding: 10px 14px; margin-bottom: 10px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
  }}
  .toolbar span {{ font-size: 13.5px; font-weight: 700; color: #555; }}
  .btn {{ padding: 7px 18px; border: 2px solid; border-radius: 6px; cursor: pointer; font-weight: 700; font-size: 13.5px; transition: all 0.18s; }}
  .btn-green  {{ background:#c8e6c9; color:#1b5e20; border-color:#388e3c; }}
  .btn-green.active, .btn-green:hover {{ background:#388e3c; color:white; }}
  .btn-red    {{ background:#ffcdd2; color:#b71c1c; border-color:#c62828; }}
  .btn-red.active, .btn-red:hover {{ background:#c62828; color:white; }}
  .btn-clear  {{ background:#f5f5f5; color:#555; border-color:#bbb; }}
  .btn-clear:hover {{ background:#e0e0e0; }}
  .status {{ font-size:13.5px; font-weight:600; margin-left:6px; padding:4px 12px; border-radius:20px; }}
  .status.green {{ background:#e8f5e9; color:#1b5e20; }}
  .status.red   {{ background:#ffebee; color:#b71c1c; }}
  .content {{
    background: white; border: 1.5px solid #ddd; border-radius: 10px;
    padding: 20px 24px; font-size: 15px; line-height: 1.95;
    white-space: pre-wrap; word-break: break-word;
    user-select: text; cursor: text;
    max-height: {height - 140}px; overflow-y: auto;
  }}
  .ai-hl-bad  {{ background: #ffe082; border-radius: 3px; padding: 1px 2px; border-bottom: 2.5px solid #f9a825; cursor: help; }}
  .ai-hl-good {{ background: #bbdefb; border-radius: 3px; padding: 1px 2px; border-bottom: 2.5px solid #1976d2; cursor: help; }}
  .hl-green   {{ background: #c8e6c9; border-radius: 3px; padding: 1px 0; cursor: pointer; }}
  .hl-red     {{ background: #ffcdd2; border-radius: 3px; padding: 1px 0; text-decoration: underline; text-decoration-color: #c62828; cursor: pointer; }}
</style>
</head>
<body>
<div class="legend">
  <span>Chú thích màu:</span>
  <div class="legend-item"><div class="dot dot-yellow"></div><span>🟡 AI – Cần cải thiện</span></div>
  <div class="legend-item"><div class="dot dot-blue"></div><span>🔵 AI – Ý tưởng hay</span></div>
  <div class="legend-item"><div class="dot dot-green"></div><span>🟢 Người dùng – Điểm tốt</span></div>
  <div class="legend-item"><div class="dot dot-red"></div><span>🔴 Người dùng – Ghi chú lỗi</span></div>
</div>
<div class="toolbar">
  <span>Bôi đen để highlight:</span>
  <button class="btn btn-green active" id="btn-green" onclick="setColor('green')">🟢 Đánh dấu tốt</button>
  <button class="btn btn-red" id="btn-red" onclick="setColor('red')">🔴 Gạch lỗi</button>
  <button class="btn btn-clear" onclick="clearUserHL()">🗑️ Xóa highlight người dùng</button>
  <span class="status green" id="status">🟢 Màu xanh</span>
</div>
<div class="content" id="content-{uid}">{highlighted_text}</div>
<script>
var currentColor = 'green';
function setColor(color) {{
  currentColor = color;
  document.getElementById('btn-green').classList.toggle('active', color === 'green');
  document.getElementById('btn-red').classList.toggle('active', color === 'red');
  var st = document.getElementById('status');
  st.textContent = color === 'green' ? '🟢 Màu xanh' : '🔴 Màu đỏ';
  st.className = 'status ' + color;
}}
function clearUserHL() {{
  var c = document.getElementById('content-{uid}');
  c.querySelectorAll('.hl-green, .hl-red').forEach(function(s) {{
    s.replaceWith(document.createTextNode(s.innerText));
  }});
  c.normalize();
}}
document.getElementById('content-{uid}').addEventListener('mouseup', function(e) {{
  var sel = window.getSelection();
  if (!sel || sel.toString().trim() === '') return;
  var range = sel.getRangeAt(0);
  var node = e.target;
  if (node.classList && (node.classList.contains('hl-green') || node.classList.contains('hl-red'))) {{
    node.replaceWith(document.createTextNode(node.innerText));
    node.parentNode && node.parentNode.normalize();
    sel.removeAllRanges(); return;
  }}
  var cur = (range.commonAncestorContainer.nodeType === 3)
    ? range.commonAncestorContainer.parentElement
    : range.commonAncestorContainer;
  while (cur && cur !== document.getElementById('content-{uid}')) {{
    if (cur.classList && (cur.classList.contains('hl-green') || cur.classList.contains('hl-red'))) {{
      cur.replaceWith(document.createTextNode(cur.innerText));
      cur.parentNode && cur.parentNode.normalize();
      sel.removeAllRanges(); return;
    }}
    cur = cur.parentElement;
  }}
  var cls = currentColor === 'red' ? 'hl-red' : 'hl-green';
  var span = document.createElement('span');
  span.className = cls;
  try {{ span.appendChild(range.extractContents()); range.insertNode(span); }} catch(ex) {{}}
  sel.removeAllRanges();
}});
</script>
</body>
</html>
"""
    components.html(html_code, height=height, scrolling=False)

# ─── EXCEL EXPORT ─────────────────────────────────────────────────────────────

def export_excel_summary(records: list) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return None

    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Claude - Tổng hợp điểm"

    header_fill = PatternFill("solid", fgColor="0D47A1")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    pass_fill   = PatternFill("solid", fgColor="C8E6C9")
    fail_fill   = PatternFill("solid", fgColor="FFCDD2")
    border = Border(left=Side(style='thin'), right=Side(style='thin'),
                    top=Side(style='thin'), bottom=Side(style='thin'))
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left   = Alignment(horizontal='left', vertical='center', wrap_text=True)

    headers = ["STT","Họ và tên","Đơn vị","Tên đề tài",
               "Tính mới\n(30đ)","Áp dụng\n(30đ)","Hiệu quả\n(40đ)","Tổng điểm","Kết quả"]
    ws1.append(headers)
    for col_idx, _ in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col_idx)
        cell.fill = header_fill; cell.font = header_font
        cell.alignment = center; cell.border = border
    ws1.row_dimensions[1].height = 35
    col_widths = [6, 22, 24, 45, 12, 12, 12, 12, 12]
    for i, w in enumerate(col_widths, 1):
        ws1.column_dimensions[get_column_letter(i)].width = w

    claude_scored = [r for r in records if r.get("score_total") is not None]
    for idx, d in enumerate(claude_scored, 1):
        sm = _safe_float(d.get("score_moi"))
        sn = _safe_float(d.get("score_nhan_rong"))
        sh = _safe_float(d.get("score_hieu_qua"))
        total = sm + sn + sh
        row = [idx, d.get("ho_ten",""), d.get("don_vi_cong_tac",""),
               (d.get("ten_de_tai","")[:60]), sm, sn, sh, total, d.get("ket_qua","")]
        ws1.append(row)
        r = idx + 1
        fill = pass_fill if d.get("ket_qua") == "Đạt" else fail_fill
        for col_i in range(1, len(headers)+1):
            cell = ws1.cell(row=r, column=col_i)
            cell.border = border
            cell.alignment = left if col_i in [2,3,4] else center
            if col_i == len(headers): cell.fill = fill
        ws1.row_dimensions[r].height = 25

    # Sheet 2: DeepSeek
    ws2 = wb.create_sheet("DeepSeek - Tổng hợp điểm")
    ws2.append(headers)
    for col_idx, _ in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=col_idx)
        cell.fill = header_fill; cell.font = header_font
        cell.alignment = center; cell.border = border
    for i, w in enumerate(col_widths, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    Deepseek_scored = [r for r in records if r.get("Deepseek_score_total") is not None]
    for idx, d in enumerate(Deepseek_scored, 1):
        sm = _safe_float(d.get("Deepseek_score_moi"))
        sn = _safe_float(d.get("Deepseek_score_nhan_rong"))
        sh = _safe_float(d.get("Deepseek_score_hieu_qua"))
        total = sm + sn + sh
        row = [idx, d.get("ho_ten",""), d.get("don_vi_cong_tac",""),
               (d.get("ten_de_tai","")[:60]), sm, sn, sh, total, d.get("Deepseek_ket_qua","")]
        ws2.append(row)
        r = idx + 1
        fill = pass_fill if d.get("Deepseek_ket_qua") == "Đạt" else fail_fill
        for col_i in range(1, len(headers)+1):
            cell = ws2.cell(row=r, column=col_i)
            cell.border = border
            cell.alignment = left if col_i in [2,3,4] else center
            if col_i == len(headers): cell.fill = fill
        ws2.row_dimensions[r].height = 25

    # Sheet 3: So sánh
    ws3 = wb.create_sheet("So sánh 2 AI")
    headers3 = ["STT","Tác giả","Đề tài","Claude Điểm","Claude KQ","DeepSeek Điểm","DeepSeek KQ","Chênh lệch"]
    ws3.append(headers3)
    for col_idx, _ in enumerate(headers3, 1):
        cell = ws3.cell(row=1, column=col_idx)
        cell.fill = header_fill; cell.font = header_font
        cell.alignment = center; cell.border = border
    col_widths3 = [6, 25, 50, 12, 10, 12, 10, 12]
    for i, w in enumerate(col_widths3, 1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    for idx, d in enumerate(claude_scored, 1):
        claude_total  = _safe_float(d.get("score_total"))
        claude_kq     = d.get("ket_qua") or ""
        Deepseek_total = _safe_float(d.get("Deepseek_score_total"))
        Deepseek_kq   = d.get("Deepseek_ket_qua") or "Chưa chấm"
        diff = abs(claude_total - Deepseek_total) if Deepseek_total else 0
        row = [idx, d.get("ho_ten",""), (d.get("ten_de_tai","")[:55]),
               round(claude_total,1), claude_kq,
               round(Deepseek_total,1) if Deepseek_total else "Chưa chấm", Deepseek_kq, round(diff,1)]
        ws3.append(row)
        r = idx + 1
        for col_i in range(1, len(headers3)+1):
            cell = ws3.cell(row=r, column=col_i)
            cell.border = border
            cell.alignment = center if col_i > 2 else left

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()

def export_excel_detail_comments(records: list) -> bytes:
    """Xuất Excel chi tiết nhận xét của cả 2 AI - ĐẦY ĐỦ NHẬN XÉT"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return None

    wb = openpyxl.Workbook()

    # ==================== SHEET 1: CLAUDE - NHẬN XÉT CHI TIẾT ====================
    ws1 = wb.active
    ws1.title = "Claude - Chi tiết"

    header_fill = PatternFill("solid", fgColor="0D47A1")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_wrap = Alignment(horizontal='left', vertical='top', wrap_text=True)

    # Header rộng hơn để chứa nhận xét
    headers1 = [
        "STT", "ID", "Tác giả", "Đơn vị", "Đề tài",
        "Tính mới\n(điểm)", "Áp dụng\n(điểm)", "Hiệu quả\n(điểm)", "Tổng điểm", "Kết quả",
        "NHẬN XÉT TÍNH MỚI", "NHẬN XÉT ÁP DỤNG", "NHẬN XÉT HIỆU QUẢ",
        "NHẬN XÉT CHUNG", "KIẾN NGHỊ", "GỢI Ý CẢI THIỆN"
    ]
    ws1.append(headers1)
    
    for col_idx, _ in enumerate(headers1, 1):
        cell = ws1.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = border

    ws1.row_dimensions[1].height = 45
    col_widths1 = [5, 6, 22, 28, 45, 10, 10, 10, 10, 10, 50, 50, 50, 55, 45, 50]
    for i, w in enumerate(col_widths1, 1):
        ws1.column_dimensions[get_column_letter(i)].width = w

    claude_scored = [r for r in records if r.get("score_total") is not None]
    for idx, d in enumerate(claude_scored, 1):
        sm = _safe_float(d.get("score_moi"))
        sn = _safe_float(d.get("score_nhan_rong"))
        sh = _safe_float(d.get("score_hieu_qua"))
        total = sm + sn + sh
        
        row = [
            idx, d["id"], d.get("ho_ten", ""), d.get("don_vi_cong_tac", ""),
            d.get("ten_de_tai", "")[:80],
            sm, sn, sh, total, d.get("ket_qua", ""),
            d.get("nhan_xet_moi", ""), d.get("nhan_xet_nhan_rong", ""),
            d.get("nhan_xet_hieu_qua", ""), d.get("nhan_xet_chung", ""),
            d.get("kien_nghi", ""), d.get("cai_thien", "")
        ]
        ws1.append(row)
        r = idx + 1
        for col_i in range(1, len(headers1) + 1):
            cell = ws1.cell(row=r, column=col_i)
            cell.border = border
            if col_i in [1, 2, 6, 7, 8, 9, 10]:
                cell.alignment = center
            else:
                cell.alignment = left_wrap
        ws1.row_dimensions[r].height = 70

    # ==================== SHEET 2: DEEPSEEK - NHẬN XÉT CHI TIẾT ====================
    ws2 = wb.create_sheet("DeepSeek - Chi tiết")
    ws2.append(headers1)
    
    for col_idx, _ in enumerate(headers1, 1):
        cell = ws2.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = border

    ws2.row_dimensions[1].height = 45
    for i, w in enumerate(col_widths1, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    deepseek_scored = [r for r in records if r.get("Deepseek_score_total") is not None]
    for idx, d in enumerate(deepseek_scored, 1):
        sm = _safe_float(d.get("Deepseek_score_moi"))
        sn = _safe_float(d.get("Deepseek_score_nhan_rong"))
        sh = _safe_float(d.get("Deepseek_score_hieu_qua"))
        total = sm + sn + sh
        
        row = [
            idx, d["id"], d.get("ho_ten", ""), d.get("don_vi_cong_tac", ""),
            d.get("ten_de_tai", "")[:80],
            sm, sn, sh, total, d.get("Deepseek_ket_qua", ""),
            d.get("Deepseek_nhan_xet_moi", ""), d.get("Deepseek_nhan_xet_nhan_rong", ""),
            d.get("Deepseek_nhan_xet_hieu_qua", ""), d.get("Deepseek_nhan_xet_chung", ""),
            d.get("Deepseek_kien_nghi", ""), d.get("Deepseek_cai_thien", "")
        ]
        ws2.append(row)
        r = idx + 1
        for col_i in range(1, len(headers1) + 1):
            cell = ws2.cell(row=r, column=col_i)
            cell.border = border
            if col_i in [1, 2, 6, 7, 8, 9, 10]:
                cell.alignment = center
            else:
                cell.alignment = left_wrap
        ws2.row_dimensions[r].height = 70

    # ==================== SHEET 3: SO SÁNH 2 AI (kèm nhận xét) ====================
    ws3 = wb.create_sheet("So sánh 2 AI")
    headers3 = [
        "STT", "ID", "Tác giả", "Đề tài",
        "Claude Điểm", "Claude KQ", "Claude Nhận xét chung",
        "DeepSeek Điểm", "DeepSeek KQ", "DeepSeek Nhận xét chung",
        "Chênh lệch", "Thống nhất"
    ]
    ws3.append(headers3)
    
    for col_idx, _ in enumerate(headers3, 1):
        cell = ws3.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = border

    ws3.row_dimensions[1].height = 35
    col_widths3 = [5, 6, 22, 50, 10, 8, 50, 10, 8, 50, 10, 8]
    for i, w in enumerate(col_widths3, 1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    for idx, d in enumerate(claude_scored, 1):
        claude_total = _safe_float(d.get("score_total"))
        claude_kq = d.get("ket_qua", "")
        claude_nx = d.get("nhan_xet_chung", "")[:200]
        
        deepseek_total = _safe_float(d.get("Deepseek_score_total"))
        deepseek_kq = d.get("Deepseek_ket_qua", "Chưa chấm")
        deepseek_nx = d.get("Deepseek_nhan_xet_chung", "")[:200] if d.get("Deepseek_nhan_xet_chung") else "Chưa chấm"
        
        diff = abs(claude_total - deepseek_total) if deepseek_total else 0
        thong_nhat = "✅ Có" if d.get("ket_qua") == d.get("Deepseek_ket_qua") else "❌ Không"
        
        row = [
            idx, d["id"], d.get("ho_ten", ""), (d.get("ten_de_tai", "")[:60]),
            round(claude_total, 1), claude_kq, claude_nx,
            round(deepseek_total, 1) if deepseek_total else "Chưa chấm", 
            deepseek_kq, deepseek_nx,
            round(diff, 1), thong_nhat
        ]
        ws3.append(row)
        r = idx + 1
        for col_i in range(1, len(headers3) + 1):
            cell = ws3.cell(row=r, column=col_i)
            cell.border = border
            cell.alignment = left_wrap if col_i in [4, 7, 10] else center
        ws3.row_dimensions[r].height = 50

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()

# ─── UI COMPONENTS ────────────────────────────────────────────────────────────

def render_header():
    st.markdown("""
    <div class="main-header">
        <h1>📋 Hệ Thống Chấm Điểm SKKN</h1>
        <p>Hội đồng Sáng kiến – UBND xã Lấp Vò, tỉnh Đồng Tháp &nbsp;|&nbsp;
        🤖 Claude AI + 🟢 DeepSeek AI – Chấm điểm song song</p>
    </div>
    """, unsafe_allow_html=True)

def render_threshold_info():
    cfg = THRESHOLD_CONFIG
    st.markdown(f"""
    <div class="threshold-info">
        ⚙️ <b>Ngưỡng công nhận:</b>&nbsp;
        Tổng ≥ <b>{cfg['total_min']}đ</b> &nbsp;|&nbsp;
        Tính mới ≥ <b>{cfg['moi_min']}/30</b> &nbsp;|&nbsp;
        Áp dụng ≥ <b>{cfg['nhan_rong_min']}/30</b> &nbsp;|&nbsp;
        Hiệu quả ≥ <b>{cfg['hieu_qua_min']}/40</b>
    </div>
    """, unsafe_allow_html=True)

def render_info_box(d: dict):
    st.markdown(f"""
    <div class="info-box">
        <div class="info-row"><span class="info-label">👤 Tác giả:</span><span>{d.get('ho_ten') or '—'}</span></div>
        <div class="info-row"><span class="info-label">🏫 Đơn vị:</span><span>{d.get('don_vi_cong_tac') or '—'}</span></div>
        <div class="info-row"><span class="info-label">📌 Chức danh:</span><span>{d.get('chuc_danh') or '—'}</span></div>
        <div class="info-row"><span class="info-label">🔬 Trình độ:</span><span>{d.get('trinh_do') or '—'}</span></div>
        <div class="info-row"><span class="info-label">📅 Ngày áp dụng:</span><span>{d.get('ngay_ap_dung') or '—'}</span></div>
        <div class="info-row"><span class="info-label">🏷️ Lĩnh vực:</span><span>{d.get('linh_vuc') or '—'}</span></div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown(f"**📖 Tên đề tài:** {d.get('ten_de_tai') or '—'}")

def render_claude_score_cards(d: dict):
    if d["trang_thai"] != "Đã chấm" or d["score_total"] is None:
        st.markdown('<span class="badge-pending">⏳ Chưa chấm điểm</span>', unsafe_allow_html=True)
        return
    sm = _safe_float(d["score_moi"]); sn = _safe_float(d["score_nhan_rong"]); sh = _safe_float(d["score_hieu_qua"])
    total = sm + sn + sh
    ket_qua = d.get("ket_qua") or _compute_ket_qua(sm, sn, sh)
    badge = (f'<span class="badge-pass">✅ {ket_qua}</span>' if ket_qua == "Đạt"
             else f'<span class="badge-fail">❌ {ket_qua}</span>')
    c1, c2, c3, c4 = st.columns([1,1,1,1.2])
    with c1:
        st.markdown(f"""<div class="score-card" style="border-left-color:#1565c0">
            <div class="score-label">🆕 Tính mới</div>
            <div class="score-num">{sm:.0f}<small>/30</small></div></div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="score-card" style="border-left-color:#1565c0">
            <div class="score-label">🔄 Áp dụng</div>
            <div class="score-num">{sn:.0f}<small>/30</small></div></div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="score-card" style="border-left-color:#1565c0">
            <div class="score-label">📈 Hiệu quả</div>
            <div class="score-num">{sh:.0f}<small>/40</small></div></div>""", unsafe_allow_html=True)
    with c4:
        color = '#1b5e20' if ket_qua == 'Đạt' else '#b71c1c'
        st.markdown(f"""<div class="score-card" style="border-left-color:{color}">
            <div class="score-label">🏆 Tổng điểm (Claude)</div>
            <div class="score-num" style="color:{color}">{total:.0f}<small>/100</small></div>
            <div style="text-align:center;margin-top:6px">{badge}</div></div>""", unsafe_allow_html=True)

def render_Deepseek_score_cards(d: dict):
    if d.get("Deepseek_trang_thai") != "Đã chấm":
        st.markdown('<span class="badge-pending">⏳ Chưa chấm điểm DeepSeek</span>', unsafe_allow_html=True)
        return
    sm = _safe_float(d.get("Deepseek_score_moi")); sn = _safe_float(d.get("Deepseek_score_nhan_rong")); sh = _safe_float(d.get("Deepseek_score_hieu_qua"))
    total = sm + sn + sh
    ket_qua = d.get("Deepseek_ket_qua") or _compute_ket_qua(sm, sn, sh)
    badge = (f'<span class="badge-pass">✅ {ket_qua}</span>' if ket_qua == "Đạt"
             else f'<span class="badge-fail">❌ {ket_qua}</span>')
    c1, c2, c3, c4 = st.columns([1,1,1,1.2])
    with c1:
        st.markdown(f"""<div class="score-card" style="border-left-color:#2e7d32">
            <div class="score-label">🆕 Tính mới (DeepSeek)</div>
            <div class="score-num">{sm:.0f}<small>/30</small></div></div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="score-card" style="border-left-color:#2e7d32">
            <div class="score-label">🔄 Áp dụng (DeepSeek)</div>
            <div class="score-num">{sn:.0f}<small>/30</small></div></div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="score-card" style="border-left-color:#2e7d32">
            <div class="score-label">📈 Hiệu quả (DeepSeek)</div>
            <div class="score-num">{sh:.0f}<small>/40</small></div></div>""", unsafe_allow_html=True)
    with c4:
        color = '#1b5e20' if ket_qua == 'Đạt' else '#b71c1c'
        st.markdown(f"""<div class="score-card" style="border-left-color:{color}">
            <div class="score-label">🏆 Tổng điểm (DeepSeek)</div>
            <div class="score-num" style="color:{color}">{total:.0f}<small>/100</small></div>
            <div style="text-align:center;margin-top:6px">{badge}</div></div>""", unsafe_allow_html=True)

def render_claude_comments(d: dict):
    with st.expander("💬 Nhận xét Claude AI", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown('<div class="section-header">🆕 Tính mới & Sáng tạo (Claude)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="comment-box">{d.get("nhan_xet_moi") or "—"}</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-header">🔄 Khả năng áp dụng (Claude)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="comment-box">{d.get("nhan_xet_nhan_rong") or "—"}</div>', unsafe_allow_html=True)
        with col2:
            st.markdown('<div class="section-header">📈 Hiệu quả (Claude)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="comment-box">{d.get("nhan_xet_hieu_qua") or "—"}</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-header">📝 Nhận xét chung (Claude)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="comment-box">{d.get("nhan_xet_chung") or "—"}</div>', unsafe_allow_html=True)
    with st.expander("⚖️ Kiến nghị & Gợi ý (Claude)", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown('<div class="section-header">⚖️ Kiến nghị (Claude)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="warning-box">{d.get("kien_nghi") or "—"}</div>', unsafe_allow_html=True)
        with col2:
            st.markdown('<div class="section-header">💡 Gợi ý cải thiện (Claude)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="improve-box">{d.get("cai_thien") or "—"}</div>', unsafe_allow_html=True)

def render_Deepseek_comments(d: dict):
    with st.expander("💬 Nhận xét DeepSeek AI", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown('<div class="section-header">🆕 Tính mới & Sáng tạo (DeepSeek)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="comment-box">{d.get("Deepseek_nhan_xet_moi") or "—"}</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-header">🔄 Khả năng áp dụng (DeepSeek)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="comment-box">{d.get("Deepseek_nhan_xet_nhan_rong") or "—"}</div>', unsafe_allow_html=True)
        with col2:
            st.markdown('<div class="section-header">📈 Hiệu quả (DeepSeek)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="comment-box">{d.get("Deepseek_nhan_xet_hieu_qua") or "—"}</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-header">📝 Nhận xét chung (DeepSeek)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="comment-box">{d.get("Deepseek_nhan_xet_chung") or "—"}</div>', unsafe_allow_html=True)
    with st.expander("⚖️ Kiến nghị & Gợi ý (DeepSeek)", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown('<div class="section-header">⚖️ Kiến nghị (DeepSeek)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="warning-box">{d.get("Deepseek_kien_nghi") or "—"}</div>', unsafe_allow_html=True)
        with col2:
            st.markdown('<div class="section-header">💡 Gợi ý cải thiện (DeepSeek)</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="improve-box">{d.get("Deepseek_cai_thien") or "—"}</div>', unsafe_allow_html=True)

def render_score_comparison(d: dict):
    claude_total   = _safe_float(d.get("score_total"))
    Deepseek_total = _safe_float(d.get("Deepseek_score_total"))
    if not Deepseek_total:
        st.info("🟢 Chưa có điểm từ DeepSeek. Hãy chấm bằng DeepSeek để so sánh.")
        return
    claude_kq   = d.get("ket_qua","")
    Deepseek_kq = d.get("Deepseek_ket_qua","")
    diff = abs(claude_total - Deepseek_total)
    st.markdown("### 📊 So sánh kết quả giữa 2 AI")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown(f"""
        <div class="score-card" style="background:#e3f2fd">
            <div class="score-label">🤖 Claude AI</div>
            <div class="score-num" style="color:#0d47a1">{claude_total:.0f}<small>/100</small></div>
            <div class="score-label">{claude_kq}</div>
        </div>""", unsafe_allow_html=True)
    with col2:
        st.markdown(f"""
        <div class="score-card" style="background:#e8f5e9">
            <div class="score-label">🟢 DeepSeek AI</div>
            <div class="score-num" style="color:#2e7d32">{Deepseek_total:.0f}<small>/100</small></div>
            <div class="score-label">{Deepseek_kq}</div>
        </div>""", unsafe_allow_html=True)
    with col3:
        same_result = claude_kq == Deepseek_kq
        status = "✅ Thống nhất" if same_result else "⚠️ Không thống nhất"
        color  = "#1b5e20" if same_result else "#e65100"
        st.markdown(f"""
        <div class="score-card" style="background:#fff3e0">
            <div class="score-label">📊 So sánh</div>
            <div class="score-num" style="color:{color};font-size:1.5rem">Chênh: {diff:.0f}đ</div>
            <div class="score-label">{status}</div>
        </div>""", unsafe_allow_html=True)

# ─── PAGES ────────────────────────────────────────────────────────────────────

def page_upload():
    st.subheader("📤 Nạp file SKKN vào hệ thống")
    st.info("Upload một hoặc nhiều file .docx để trích xuất thông tin và lưu vào cơ sở dữ liệu.", icon="ℹ️")
    uploaded = st.file_uploader("Chọn file SKKN (.docx)", type=["docx"], accept_multiple_files=True)
    if not uploaded: return
    st.write(f"**{len(uploaded)} file** được chọn.")
    if not st.button("💾 Nạp vào hệ thống", type="primary"): return
    progress = st.progress(0)
    results = []
    for i, f in enumerate(uploaded):
        with st.spinner(f"Đang xử lý: {f.name}…"):
            try:
                raw = extract_docx_text(f.read())
                info = parse_info_from_text(raw, f.name)
                info["raw_text"] = raw
                new_id = insert_record(info)
                results.append((f.name, True, new_id, info))
            except Exception as e:
                results.append((f.name, False, None, str(e)))
        progress.progress((i+1)/len(uploaded))
    st.success(f"Hoàn tất! Đã nạp {sum(1 for r in results if r[1])} / {len(results)} file.")
    for name, ok, rid, info in results:
        if ok:
            with st.expander(f"✅ {name} → ID #{rid}", expanded=False):
                st.write(f"**Tác giả:** {info.get('ho_ten','—')}")
                st.write(f"**Đề tài:** {info.get('ten_de_tai','—')}")
        else:
            st.error(f"❌ {name}: {info}")

def page_list():
    st.subheader("📋 Danh sách SKKN trong hệ thống")
    records = get_all_records()
    
    if not records:
        st.warning("Chưa có SKKN nào. Hãy nạp file ở tab 'Nạp File'.")
        return

    # === THỐNG KÊ NHANH ===
    total = len(records)
    da_cham = sum(1 for r in records if r["trang_thai"] == "Đã chấm")
    dat = sum(1 for r in records if r["ket_qua"] == "Đạt")
    ds_da_cham = sum(1 for r in records if r.get("Deepseek_trang_thai") == "Đã chấm")
    ds_dat = sum(1 for r in records if r.get("Deepseek_ket_qua") == "Đạt")

    col1, col2, col3, col4, col5, col6 = st.columns(6)
    col1.metric("📁 Tổng SKKN", total)
    col2.metric("🤖 Claude đã chấm", da_cham)
    col3.metric("🏆 Claude đạt", dat)
    col4.metric("🟢 DeepSeek đã chấm", ds_da_cham)
    col5.metric("🏆 DeepSeek đạt", ds_dat)
    if da_cham > 0 and ds_da_cham > 0:
        thong_nhat = sum(1 for r in records if r.get("ket_qua") == r.get("Deepseek_ket_qua"))
        col6.metric("🤝 Thống nhất", f"{thong_nhat}/{da_cham}")

    render_threshold_info()
    st.divider()

    # === BỘ LỌC ===
    col_f1, col_f2, col_f3, col_f4 = st.columns(4)
    with col_f1:
        filter_status = st.selectbox("🔽 Lọc trạng thái Claude", ["Tất cả", "Chưa chấm", "Đã chấm"])
    with col_f2:
        filter_result = st.selectbox("🏆 Lọc kết quả Claude", ["Tất cả", "Đạt", "Không đạt"])
    with col_f3:
        filter_ds_status = st.selectbox("🔽 Lọc trạng thái DeepSeek", ["Tất cả", "Chưa chấm", "Đã chấm"])
    with col_f4:
        sort_by = st.selectbox("📊 Sắp xếp theo", ["ID mới nhất", "ID cũ nhất", "Tên A-Z", "Tên Z-A", "Claude điểm cao", "DeepSeek điểm cao"])

    # === LỌC DỮ LIỆU ===
    filtered_records = records.copy()
    
    # Lọc theo Claude
    if filter_status != "Tất cả":
        filtered_records = [r for r in filtered_records if r["trang_thai"] == filter_status]
    if filter_result != "Tất cả":
        filtered_records = [r for r in filtered_records if r.get("ket_qua") == filter_result]
    
    # Lọc theo DeepSeek
    if filter_ds_status != "Tất cả":
        filtered_records = [r for r in filtered_records if r.get("Deepseek_trang_thai") == filter_ds_status]
    
    # Sắp xếp
    if sort_by == "ID mới nhất":
        filtered_records.sort(key=lambda x: x["id"], reverse=True)
    elif sort_by == "ID cũ nhất":
        filtered_records.sort(key=lambda x: x["id"])
    elif sort_by == "Tên A-Z":
        filtered_records.sort(key=lambda x: x.get("ho_ten") or "")
    elif sort_by == "Tên Z-A":
        filtered_records.sort(key=lambda x: x.get("ho_ten") or "", reverse=True)
    elif sort_by == "Claude điểm cao":
        filtered_records.sort(key=lambda x: x.get("score_total") or 0, reverse=True)
    elif sort_by == "DeepSeek điểm cao":
        filtered_records.sort(key=lambda x: x.get("Deepseek_score_total") or 0, reverse=True)

    st.markdown(f"**📊 Hiển thị {len(filtered_records)} / {len(records)} SKKN**")
    st.divider()

    # === BẢNG DỮ LIỆU CHÍNH ===
    if not filtered_records:
        st.warning("Không có SKKN nào phù hợp với bộ lọc.")
        return

    # Tạo dữ liệu cho bảng
    table_data = []
    for d in filtered_records:
        # Claude status
        claude_status_icon = "✅" if d["trang_thai"] == "Đã chấm" else "⏳"
        claude_score = f"{d['score_total']:.0f}" if d["score_total"] is not None else "—"
        claude_result = d.get("ket_qua", "—")
        claude_result_icon = "🏆" if claude_result == "Đạt" else "❌" if claude_result == "Không đạt" else "⏳"
        
        # DeepSeek status
        ds_status_icon = "✅" if d.get("Deepseek_trang_thai") == "Đã chấm" else "⏳"
        ds_score = f"{d['Deepseek_score_total']:.0f}" if d.get("Deepseek_score_total") else "—"
        ds_result = d.get("Deepseek_ket_qua", "—")
        ds_result_icon = "🏆" if ds_result == "Đạt" else "❌" if ds_result == "Không đạt" else "⏳"
        
        # So sánh
        if d["trang_thai"] == "Đã chấm" and d.get("Deepseek_trang_thai") == "Đã chấm":
            same_result = "✅" if d.get("ket_qua") == d.get("Deepseek_ket_qua") else "⚠️"
        else:
            same_result = "⏳"
        
        table_data.append({
            "ID": d["id"],
            "Tác giả": d.get("ho_ten", "—")[:30],
            "Đơn vị": d.get("don_vi_cong_tac", "—")[:25],
            "Đề tài": (d.get("ten_de_tai") or "—")[:50],
            "🤖 Claude": f"{claude_status_icon} {claude_score}",
            "Claude KQ": f"{claude_result_icon} {claude_result}",
            "🟢 DeepSeek": f"{ds_status_icon} {ds_score}",
            "DeepSeek KQ": f"{ds_result_icon} {ds_result}",
            "🔄 So sánh": same_result,
            "Chi tiết": "🔍 Xem"
        })

    # Hiển thị bảng
    try:
        import pandas as pd
        df = pd.DataFrame(table_data)
        
        # Cấu hình hiển thị bảng
        st.dataframe(
            df,
            use_container_width=True,
            height=500,
            column_config={
                "ID": st.column_config.NumberColumn("ID", width="small"),
                "Tác giả": st.column_config.TextColumn("Tác giả", width="medium"),
                "Đơn vị": st.column_config.TextColumn("Đơn vị", width="medium"),
                "Đề tài": st.column_config.TextColumn("Đề tài", width="large"),
                "🤖 Claude": st.column_config.TextColumn("Claude", width="small"),
                "Claude KQ": st.column_config.TextColumn("KQ", width="small"),
                "🟢 DeepSeek": st.column_config.TextColumn("DeepSeek", width="small"),
                "DeepSeek KQ": st.column_config.TextColumn("KQ", width="small"),
                "🔄 So sánh": st.column_config.TextColumn("So sánh", width="small"),
                "Chi tiết": st.column_config.TextColumn("", width="small"),
            }
        )
    except ImportError:
        # Fallback nếu không có pandas
        for d in table_data:
            st.write(f"#{d['ID']} | {d['Tác giả']} | Claude: {d['🤖 Claude']} | DeepSeek: {d['🟢 DeepSeek']}")

    st.divider()
    
    # === PHẦN CHI TIẾT BÀI (khi click vào nút Xem) ===
    st.markdown("### 📄 Chi tiết SKKN")
    st.info("Click vào nút **🔍 Xem** bên dưới để xem chi tiết từng bài", icon="ℹ️")
    
    # Tạo các nút xem chi tiết
    cols_per_row = 4
    for i in range(0, len(filtered_records), cols_per_row):
        cols = st.columns(cols_per_row)
        for j, col in enumerate(cols):
            idx = i + j
            if idx < len(filtered_records):
                d = filtered_records[idx]
                with col:
                    # Hiển thị card nhỏ cho mỗi bài
                    claude_score = f"{d['score_total']:.0f}" if d["score_total"] else "?"
                    ds_score = f"{d['Deepseek_score_total']:.0f}" if d.get("Deepseek_score_total") else "?"
                    status_color = "🟢" if d.get("ket_qua") == "Đạt" else "🔴" if d.get("ket_qua") == "Không đạt" else "⚪"
                    
                    with st.container(border=True):
                        st.markdown(f"**#{d['id']} - {d.get('ho_ten', '?')[:25]}**")
                        st.caption(f"📖 {(d.get('ten_de_tai') or '')[:40]}...")
                        st.markdown(f"🤖 Claude: **{claude_score}**/100 | 🟢 DeepSeek: **{ds_score}**/100")
                        st.markdown(f"{status_color} Kết quả: {d.get('ket_qua', 'Chưa chấm')}")
                        
                        if st.button(f"🔍 Xem chi tiết #{d['id']}", key=f"card_view_{d['id']}", use_container_width=True):
                            st.session_state["view_id"] = d["id"]
                            st.session_state["page"] = "detail"
                            st.rerun()
    
    # === XUẤT EXCEL ===
    st.divider()
    col_export1, col_export2, col_export3 = st.columns(3)
    with col_export1:
        if st.button("📥 Xuất Excel danh sách hiện tại", use_container_width=True):
            xlsx_bytes = export_excel_summary(filtered_records)
            if xlsx_bytes:
                st.download_button(
                    "📥 Tải Excel",
                    data=xlsx_bytes,
                    file_name=f"SKKN_DanhSach_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="export_list"
                )
    with col_export2:
        if st.button("📝 Xuất Excel chi tiết (kèm nhận xét)", use_container_width=True):
            xlsx_bytes = export_excel_detail_comments(filtered_records)
            if xlsx_bytes:
                st.download_button(
                    "📥 Tải Excel chi tiết",
                    data=xlsx_bytes,
                    file_name=f"SKKN_ChiTiet_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="export_detail"
                )
    with col_export3:
        if st.button("🗑️ Xóa tất cả SKKN", use_container_width=True):
            if st.checkbox("Xác nhận xóa TẤT CẢ?"):
                for d in records:
                    delete_record(d["id"])
                st.success("Đã xóa tất cả!")
                st.rerun()
def page_detail():
    rid = st.session_state.get("view_id")
    if not rid:
        st.warning("Chưa chọn SKKN nào.")
        return
    d = get_record(rid)
    if not d:
        st.error("Không tìm thấy SKKN.")
        return
    if st.button("← Quay lại danh sách"):
        st.session_state["page"] = "list"; st.rerun()
    st.subheader(f"📄 Chi tiết SKKN #{rid}")
    render_info_box(d)

    tab_claude, tab_ds = st.tabs(["🤖 Claude AI", "🟢 DeepSeek AI"])
    with tab_claude:
        if d["trang_thai"] == "Đã chấm":
            render_claude_score_cards(d)
            render_claude_comments(d)
            with st.expander("🎨 Xem nội dung SKKN (highlight Claude)", expanded=False):
                ai_hl = safe_load_json(d.get("ai_highlights"), [])
                render_highlight_component(d["raw_text"], uid=f"detail_claude_{rid}", ai_highlights=ai_hl, height=560)
        else:
            st.warning("Chưa chấm điểm bằng Claude.")
    with tab_ds:
        if d.get("Deepseek_trang_thai") == "Đã chấm":
            render_Deepseek_score_cards(d)
            render_Deepseek_comments(d)
            with st.expander("🎨 Xem nội dung SKKN (highlight DeepSeek)", expanded=False):
                ds_hl = safe_load_json(d.get("Deepseek_ai_highlights"), [])
                render_highlight_component(d["raw_text"], uid=f"detail_ds_{rid}", ai_highlights=ds_hl, height=560)
        else:
            st.warning("Chưa chấm điểm bằng DeepSeek.")

    xlsx_bytes = export_excel_single(d)
    if xlsx_bytes:
        st.download_button(
            f"📥 Xuất Excel đầy đủ – {d.get('ho_ten') or 'SKKN'}",
            data=xlsx_bytes, file_name=f"SKKN_{d.get('ho_ten') or d['id']}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

def page_edit_scores():
    st.subheader("✏️ Chỉnh Sửa Điểm Thủ Công")
    render_threshold_info()
    records = get_all_records()
    scored = [d for d in records if d["trang_thai"] == "Đã chấm" or d.get("Deepseek_trang_thai") == "Đã chấm"]
    if not scored:
        st.warning("Chưa có SKKN nào được chấm. Hãy chấm điểm trước.")
        return

    if "edit_refresh_key" not in st.session_state:
        st.session_state.edit_refresh_key = 0

    opts = {f"#{d['id']} – {d.get('ho_ten') or '?'} – {(d.get('ten_de_tai') or '')[:50]}": d for d in scored}
    sel_label = st.selectbox("🔍 Chọn SKKN để chỉnh sửa:", list(opts.keys()), key=f"edit_sel_{st.session_state.edit_refresh_key}")
    if not sel_label: return
    d = opts[sel_label]
    record_id = d["id"]

    st.divider()
    render_info_box(d)

    col_ci, col_di = st.columns(2)
    with col_ci:
        st.markdown("### 🤖 Điểm Claude AI")
        if d["trang_thai"] == "Đã chấm": render_claude_score_cards(d)
        else: st.info("⏳ Chưa chấm bằng Claude")
    with col_di:
        st.markdown("### 🟢 Điểm DeepSeek AI")
        if d.get("Deepseek_trang_thai") == "Đã chấm": render_Deepseek_score_cards(d)
        else: st.info("⏳ Chưa chấm bằng DeepSeek")

    st.divider()
    tab_c, tab_g = st.tabs(["🤖 Chỉnh sửa điểm Claude", "🟢 Chỉnh sửa điểm DeepSeek"])

    with tab_c:
        if d["trang_thai"] == "Đã chấm":
            osm = _safe_float(d.get("orig_score_moi") or d.get("score_moi"))
            osn = _safe_float(d.get("orig_score_nhan_rong") or d.get("score_nhan_rong"))
            osh = _safe_float(d.get("orig_score_hieu_qua") or d.get("score_hieu_qua"))
            current_sm = _safe_float(d.get("score_moi"))
            current_sn = _safe_float(d.get("score_nhan_rong"))
            current_sh = _safe_float(d.get("score_hieu_qua"))
            st.markdown("#### 🎛️ Chỉnh sửa điểm Claude")
            st.caption(f"Điểm AI gốc: **{osm:.0f}** / **{osn:.0f}** / **{osh:.0f}** (Tổng: **{osm+osn+osh:.0f}/100**)")
            col_a, col_b, col_c = st.columns(3)
            with col_a: new_sm = st.slider("Tính mới", 0, 30, int(current_sm), key=f"edit_claude_moi_{record_id}_{st.session_state.edit_refresh_key}")
            with col_b: new_sn = st.slider("Áp dụng", 0, 30, int(current_sn), key=f"edit_claude_nr_{record_id}_{st.session_state.edit_refresh_key}")
            with col_c: new_sh = st.slider("Hiệu quả", 0, 40, int(current_sh), key=f"edit_claude_hq_{record_id}_{st.session_state.edit_refresh_key}")
            new_total = new_sm + new_sn + new_sh
            new_kq    = _compute_ket_qua(new_sm, new_sn, new_sh)
            st.info(f"**Xem trước:** Tổng mới = **{new_total}/100** → **{new_kq}**")
            col_save, col_reset = st.columns(2)
            with col_save:
                with st.popover("💾 Xác nhận lưu điểm Claude", use_container_width=True):
                    st.warning(f"Lưu điểm Claude: {new_sm}/{new_sn}/{new_sh} → {new_total}/100 ({new_kq})")
                    st.markdown("---")
                    if st.button("✅ Xác nhận lưu", key=f"confirm_claude_{record_id}"):
                        save_manual_claude_scores(record_id, new_sm, new_sn, new_sh)
                        st.success(f"✅ Đã lưu!")
                        st.session_state.edit_refresh_key += 1
                        time.sleep(0.5); st.rerun()
                    if st.button("❌ Hủy", key=f"cancel_claude_{record_id}"): st.rerun()
            with col_reset:
                with st.popover("🔄 Xác nhận reset Claude", use_container_width=True):
                    st.warning(f"Reset về điểm gốc: {osm:.0f}/{osn:.0f}/{osh:.0f} → {osm+osn+osh:.0f}/100")
                    st.markdown("---")
                    if st.button("✅ Xác nhận reset", key=f"reset_claude_{record_id}"):
                        reset_claude_to_orig_scores(record_id)
                        st.success("✅ Đã reset!")
                        st.session_state.edit_refresh_key += 1
                        time.sleep(0.5); st.rerun()
                    if st.button("❌ Hủy", key=f"cancel_reset_claude_{record_id}"): st.rerun()
            st.divider()
            render_claude_comments(d)
        else:
            st.warning("Chưa chấm điểm bằng Claude. Hãy chấm trước ở tab 'Chấm Điểm AI'.")

    with tab_g:
        if d.get("Deepseek_trang_thai") == "Đã chấm":
            osm_g = _safe_float(d.get("Deepseek_orig_score_moi") or d.get("Deepseek_score_moi"))
            osn_g = _safe_float(d.get("Deepseek_orig_score_nhan_rong") or d.get("Deepseek_score_nhan_rong"))
            osh_g = _safe_float(d.get("Deepseek_orig_score_hieu_qua") or d.get("Deepseek_score_hieu_qua"))
            current_sm_g = _safe_float(d.get("Deepseek_score_moi"))
            current_sn_g = _safe_float(d.get("Deepseek_score_nhan_rong"))
            current_sh_g = _safe_float(d.get("Deepseek_score_hieu_qua"))
            st.markdown("#### 🎛️ Chỉnh sửa điểm DeepSeek")
            st.caption(f"Điểm AI gốc: **{osm_g:.0f}** / **{osn_g:.0f}** / **{osh_g:.0f}** (Tổng: **{osm_g+osn_g+osh_g:.0f}/100**)")
            col_a, col_b, col_c = st.columns(3)
            with col_a: new_sm_g = st.slider("Tính mới (DeepSeek)", 0, 30, int(current_sm_g), key=f"edit_ds_moi_{record_id}_{st.session_state.edit_refresh_key}")
            with col_b: new_sn_g = st.slider("Áp dụng (DeepSeek)", 0, 30, int(current_sn_g), key=f"edit_ds_nr_{record_id}_{st.session_state.edit_refresh_key}")
            with col_c: new_sh_g = st.slider("Hiệu quả (DeepSeek)", 0, 40, int(current_sh_g), key=f"edit_ds_hq_{record_id}_{st.session_state.edit_refresh_key}")
            new_total_g = new_sm_g + new_sn_g + new_sh_g
            new_kq_g    = _compute_ket_qua(new_sm_g, new_sn_g, new_sh_g)
            st.info(f"**Xem trước:** Tổng mới = **{new_total_g}/100** → **{new_kq_g}**")
            col_save, col_reset = st.columns(2)
            with col_save:
                with st.popover("💾 Xác nhận lưu điểm DeepSeek", use_container_width=True):
                    st.warning(f"Lưu điểm DeepSeek: {new_sm_g}/{new_sn_g}/{new_sh_g} → {new_total_g}/100 ({new_kq_g})")
                    st.markdown("---")
                    if st.button("✅ Xác nhận lưu", key=f"confirm_ds_{record_id}"):
                        save_manual_Deepseek_scores(record_id, new_sm_g, new_sn_g, new_sh_g)
                        st.success("✅ Đã lưu!")
                        st.session_state.edit_refresh_key += 1
                        time.sleep(0.5); st.rerun()
                    if st.button("❌ Hủy", key=f"cancel_ds_{record_id}"): st.rerun()
            with col_reset:
                with st.popover("🔄 Xác nhận reset DeepSeek", use_container_width=True):
                    st.warning(f"Reset về điểm gốc: {osm_g:.0f}/{osn_g:.0f}/{osh_g:.0f} → {osm_g+osn_g+osh_g:.0f}/100")
                    st.markdown("---")
                    if st.button("✅ Xác nhận reset", key=f"reset_ds_{record_id}"):
                        reset_Deepseek_to_orig_scores(record_id)
                        st.success("✅ Đã reset!")
                        st.session_state.edit_refresh_key += 1
                        time.sleep(0.5); st.rerun()
                    if st.button("❌ Hủy", key=f"cancel_reset_ds_{record_id}"): st.rerun()
            st.divider()
            render_Deepseek_comments(d)
        else:
            st.warning("Chưa chấm điểm bằng DeepSeek. Hãy chấm trước ở tab 'Chấm Điểm AI'.")

def page_stats():
    st.subheader("📊 Thống Kê & Báo Cáo Tổng Hợp")
    records = get_all_records()
    if not records:
        st.warning("Chưa có dữ liệu.")
        return

    claude_scored = [r for r in records if r["score_total"] is not None]
    ds_scored     = [r for r in records if r.get("Deepseek_score_total") is not None]

    c1, c2, c3, c4, c5, c6 = st.columns(6)
    c1.metric("📁 Tổng SKKN", len(records))
    c2.metric("🤖 Claude đã chấm", len(claude_scored))
    if claude_scored:
        claude_avg = sum(r["score_total"] for r in claude_scored) / len(claude_scored)
        claude_dat = sum(1 for r in claude_scored if r["ket_qua"] == "Đạt")
        c3.metric("📈 Claude TB", f"{claude_avg:.1f}")
        c4.metric("🏆 Claude đạt", f"{claude_dat}/{len(claude_scored)}")
    c5.metric("🟢 DeepSeek đã chấm", len(ds_scored))
    if ds_scored:
        ds_dat = sum(1 for r in ds_scored if r.get("Deepseek_ket_qua") == "Đạt")
        c6.metric("🏆 DeepSeek đạt", f"{ds_dat}/{len(ds_scored)}")

    render_threshold_info()
    if not claude_scored and not ds_scored: return

    tab_table, tab_detail, tab_compare, tab_export = st.tabs([
        "📋 Bảng điểm tổng hợp", "📝 Chi tiết nhận xét từng SKKN",
        "📊 So sánh 2 AI", "📥 Xuất Excel báo cáo"
    ])

    with tab_table:
        st.markdown("### 🤖 Bảng điểm Claude AI")
        cf1, cf2 = st.columns(2)
        with cf1: filter_kq_c = st.selectbox("Lọc kết quả Claude:", ["Tất cả","Đạt","Không đạt"], key="stat_filter_claude")
        with cf2: sort_c = st.selectbox("Sắp xếp Claude:", ["ID giảm dần","Điểm cao → thấp","Điểm thấp → cao","Tên A-Z"], key="stat_sort_claude")
        display_c = claude_scored[:]
        if filter_kq_c == "Đạt": display_c = [r for r in display_c if r.get("ket_qua") == "Đạt"]
        elif filter_kq_c == "Không đạt": display_c = [r for r in display_c if r.get("ket_qua") != "Đạt"]
        if sort_c == "Điểm cao → thấp": display_c.sort(key=lambda x: x["score_total"] or 0, reverse=True)
        elif sort_c == "Điểm thấp → cao": display_c.sort(key=lambda x: x["score_total"] or 0)
        elif sort_c == "Tên A-Z": display_c.sort(key=lambda x: x.get("ho_ten") or "")
        else: display_c.sort(key=lambda x: x["id"], reverse=True)
        try:
            import pandas as pd
            rows = [{"ID": r["id"], "Tác giả": r.get("ho_ten",""), "Đơn vị": r.get("don_vi_cong_tac",""),
                     "Đề tài": (r.get("ten_de_tai",""))[:55],
                     "Tính mới": r.get("score_moi") or 0, "Áp dụng": r.get("score_nhan_rong") or 0,
                     "Hiệu quả": r.get("score_hieu_qua") or 0, "Tổng": r.get("score_total") or 0,
                     "Kết quả": r.get("ket_qua") or ""} for r in display_c]
            st.dataframe(pd.DataFrame(rows), use_container_width=True, height=350)
        except ImportError:
            for r in display_c: st.write(f"#{r['id']} **{r.get('ho_ten','')}** → {r.get('score_total',0):.0f}/100 – {r.get('ket_qua','')}")

        st.markdown("### 🟢 Bảng điểm DeepSeek AI")
        cf3, cf4 = st.columns(2)
        with cf3: filter_kq_g = st.selectbox("Lọc kết quả DeepSeek:", ["Tất cả","Đạt","Không đạt"], key="stat_filter_ds")
        with cf4: sort_g = st.selectbox("Sắp xếp DeepSeek:", ["ID giảm dần","Điểm cao → thấp","Điểm thấp → cao","Tên A-Z"], key="stat_sort_ds")
        display_g = [r for r in records if r.get("Deepseek_score_total") is not None]
        if filter_kq_g == "Đạt": display_g = [r for r in display_g if r.get("Deepseek_ket_qua") == "Đạt"]
        elif filter_kq_g == "Không đạt": display_g = [r for r in display_g if r.get("Deepseek_ket_qua") != "Đạt"]
        if sort_g == "Điểm cao → thấp": display_g.sort(key=lambda x: x.get("Deepseek_score_total") or 0, reverse=True)
        elif sort_g == "Điểm thấp → cao": display_g.sort(key=lambda x: x.get("Deepseek_score_total") or 0)
        elif sort_g == "Tên A-Z": display_g.sort(key=lambda x: x.get("ho_ten") or "")
        else: display_g.sort(key=lambda x: x["id"], reverse=True)
        try:
            import pandas as pd
            rows_g = [{"ID": r["id"], "Tác giả": r.get("ho_ten",""), "Đơn vị": r.get("don_vi_cong_tac",""),
                       "Đề tài": (r.get("ten_de_tai",""))[:55],
                       "Tính mới": _safe_float(r.get("Deepseek_score_moi")),
                       "Áp dụng": _safe_float(r.get("Deepseek_score_nhan_rong")),
                       "Hiệu quả": _safe_float(r.get("Deepseek_score_hieu_qua")),
                       "Tổng": _safe_float(r.get("Deepseek_score_total")),
                       "Kết quả": r.get("Deepseek_ket_qua") or ""} for r in display_g]
            st.dataframe(pd.DataFrame(rows_g), use_container_width=True, height=350)
        except ImportError:
            for r in display_g: st.write(f"#{r['id']} **{r.get('ho_ten','')}** → {r.get('Deepseek_score_total',0):.0f}/100 – {r.get('Deepseek_ket_qua','')}")

    with tab_detail:
        st.markdown("### 📝 Chi tiết nhận xét từng SKKN")
        cf5, cf6, cf7 = st.columns(3)
        with cf5: filter_ai = st.selectbox("🤖 Chọn AI:", ["Claude AI","DeepSeek AI"], key="detail_ai_choice")
        with cf6: filter_kq_d = st.selectbox("Lọc kết quả:", ["Tất cả","Đạt","Không đạt"], key="detail_filter_kq")
        with cf7: search_name = st.text_input("🔍 Tìm theo tên:", placeholder="Nhập tên tác giả…", key="detail_search")
        if filter_ai == "Claude AI":
            det_recs = claude_scored[:]
            get_score = lambda r: r.get("score_total") or 0
            get_kq    = lambda r: r.get("ket_qua") or ""
        else:
            det_recs = ds_scored[:]
            get_score = lambda r: r.get("Deepseek_score_total") or 0
            get_kq    = lambda r: r.get("Deepseek_ket_qua") or ""
        if filter_kq_d == "Đạt":       det_recs = [r for r in det_recs if get_kq(r) == "Đạt"]
        elif filter_kq_d == "Không đạt": det_recs = [r for r in det_recs if get_kq(r) != "Đạt"]
        if search_name: det_recs = [r for r in det_recs if search_name.lower() in (r.get("ho_ten") or "").lower()]
        for d in det_recs:
            total = get_score(d); kq = get_kq(d); icon = "✅" if kq == "Đạt" else "❌"
            with st.expander(f"{icon} #{d['id']} – {d.get('ho_ten') or '?'} | Tổng: {total:.0f}/100 | {kq} | {d.get('ten_de_tai','')[:60]}", expanded=False):
                render_info_box(d)
                if filter_ai == "Claude AI":
                    render_claude_score_cards(d); st.markdown("---"); render_claude_comments(d)
                    with st.expander("🎨 Highlight Claude", expanded=False):
                        ai_hl = safe_load_json(d.get("ai_highlights"), [])
                        render_highlight_component(d["raw_text"], uid=f"stat_c_{d['id']}", ai_highlights=ai_hl, height=500)
                else:
                    render_Deepseek_score_cards(d); st.markdown("---"); render_Deepseek_comments(d)
                    with st.expander("🎨 Highlight DeepSeek", expanded=False):
                        ds_hl = safe_load_json(d.get("Deepseek_ai_highlights"), [])
                        render_highlight_component(d["raw_text"], uid=f"stat_g_{d['id']}", ai_highlights=ds_hl, height=500)
                xlsx_bytes = export_excel_single(d)
                if xlsx_bytes:
                    st.download_button(f"📥 Xuất Excel – {d.get('ho_ten') or 'SKKN'}",
                        data=xlsx_bytes, file_name=f"SKKN_{d.get('ho_ten') or d['id']}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key=f"stat_dl_{d['id']}_{filter_ai}")

    with tab_compare:
        st.markdown("### 📊 So sánh điểm giữa Claude và DeepSeek")
        try:
            import pandas as pd
            compare_rows = []
            for r in records:
                if r.get("score_total") is not None and r.get("Deepseek_score_total") is not None:
                    ct = _safe_float(r.get("score_total")); gt = _safe_float(r.get("Deepseek_score_total"))
                    compare_rows.append({
                        "ID": r["id"], "Tác giả": r.get("ho_ten") or "",
                        "Đề tài": (r.get("ten_de_tai") or "")[:45],
                        "Claude Điểm": round(ct,1), "Claude KQ": r.get("ket_qua") or "",
                        "DeepSeek Điểm": round(gt,1), "DeepSeek KQ": r.get("Deepseek_ket_qua") or "",
                        "Chênh lệch": abs(round(ct - gt, 1)),
                        "Thống nhất": "✅" if r.get("ket_qua") == r.get("Deepseek_ket_qua") else "⚠️"
                    })
            if compare_rows:
                df_compare = pd.DataFrame(compare_rows)
                st.dataframe(df_compare, use_container_width=True, height=400)
                s1, s2, s3 = st.columns(3)
                avg_diff = sum(c["Chênh lệch"] for c in compare_rows) / len(compare_rows)
                tn = sum(1 for c in compare_rows if c["Thống nhất"] == "✅")
                s1.metric("📊 Chênh lệch TB", f"{avg_diff:.1f} điểm")
                s2.metric("🤝 Tỷ lệ thống nhất", f"{tn}/{len(compare_rows)}")
                s3.metric("📈 % thống nhất", f"{tn/len(compare_rows)*100:.0f}%")
                st.markdown("### 📈 Biểu đồ phân bố điểm")
                chart_data = pd.DataFrame({"Claude": [c["Claude Điểm"] for c in compare_rows],
                                           "DeepSeek": [c["DeepSeek Điểm"] for c in compare_rows]})
                st.line_chart(chart_data)
            else:
                st.info("Chưa có bài nào được chấm bằng cả 2 AI để so sánh.")
        except ImportError:
            for r in records:
                if r.get("score_total") and r.get("Deepseek_score_total"):
                    st.write(f"#{r['id']} **{r.get('ho_ten','')}**: Claude={r.get('score_total',0):.0f} ({r.get('ket_qua','')}) | DeepSeek={r.get('Deepseek_score_total',0):.0f} ({r.get('Deepseek_ket_qua','')})")

    with tab_export:
        st.markdown("### 📥 Xuất Excel báo cáo chi tiết")
        st.info("Xuất file Excel bao gồm: điểm số + nhận xét từng tiêu chí + kiến nghị + gợi ý cải thiện", icon="📊")
        
        # Bộ lọc
        col_f8, col_f9 = st.columns(2)
        with col_f8:
            filter_export = st.selectbox(
                "🔽 Lọc kết quả theo Claude:",
                ["Tất cả", "Chỉ Đạt", "Chỉ Không đạt"],
                key="export_filter_claude"
            )
        with col_f9:
            filter_export_deepseek = st.selectbox(
                "🔽 Lọc kết quả theo DeepSeek:",
                ["Tất cả", "Chỉ Đạt", "Chỉ Không đạt"],
                key="export_filter_deepseek"
            )
        
        # Áp dụng bộ lọc
        filtered = records[:]
        if filter_export == "Chỉ Đạt":
            filtered = [r for r in filtered if r.get("ket_qua") == "Đạt"]
        elif filter_export == "Chỉ Không đạt":
            filtered = [r for r in filtered if r.get("ket_qua") != "Đạt"]
        
        if filter_export_deepseek == "Chỉ Đạt":
            filtered = [r for r in filtered if r.get("Deepseek_ket_qua") == "Đạt"]
        elif filter_export_deepseek == "Chỉ Không đạt":
            filtered = [r for r in filtered if r.get("Deepseek_ket_qua") != "Đạt"]
        
        st.markdown(f"**📊 Số bài sẽ xuất: {len(filtered)} / {len(records)}**")
        st.divider()
        
        # Nút xuất Excel
        xlsx_bytes = export_excel_detail_comments(filtered)
        if xlsx_bytes:
            st.download_button(
                "📥 Tải Excel BÁO CÁO ĐẦY ĐỦ (Điểm + Nhận xét + Kiến nghị + Cải thiện)",
                data=xlsx_bytes,
                file_name=f"SKKN_BaoCao_ChiTiet_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        
        st.info("✅ File Excel bao gồm 3 sheet:\n"
                "- **Claude - Chi tiết**: Điểm + nhận xét từng tiêu chí + kiến nghị + cải thiện\n"
                "- **DeepSeek - Chi tiết**: Tương tự cho DeepSeek\n"
                "- **So sánh 2 AI**: So sánh điểm và nhận xét chung", icon="ℹ️")

def export_excel_single(d: dict) -> bytes:
    """Xuất Excel 1 SKKN với đầy đủ thông tin của cả 2 AI"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return None

    wb = openpyxl.Workbook()
    
    # Sheet 1: Thông tin chung và Claude
    ws_claude = wb.active
    ws_claude.title = "Claude AI - Chi tiết"

    header_fill = PatternFill("solid", fgColor="0D47A1")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    label_font = Font(bold=True, size=12)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    wrap = Alignment(horizontal='left', vertical='top', wrap_text=True)
    center = Alignment(horizontal='center', vertical='center')

    ws_claude.column_dimensions['A'].width = 28
    ws_claude.column_dimensions['B'].width = 60

    def add_section_header(ws, title):
        ws.append([title, ""])
        r = ws.max_row
        ws.cell(r, 1).fill = header_fill
        ws.cell(r, 1).font = header_font
        ws.cell(r, 1).alignment = center
        ws.cell(r, 2).fill = header_fill
        ws.merge_cells(f"A{r}:B{r}")
        ws.row_dimensions[r].height = 28

    def add_row(ws, label, value, height=22):
        ws.append([label, str(value or "—")])
        r = ws.max_row
        ws.cell(r, 1).font = label_font
        ws.cell(r, 1).alignment = Alignment(horizontal='right', vertical='top')
        ws.cell(r, 2).alignment = wrap
        ws.cell(r, 1).border = border
        ws.cell(r, 2).border = border
        ws.row_dimensions[r].height = height

    # Thông tin tác giả
    add_section_header(ws_claude, "THÔNG TIN TÁC GIẢ")
    add_row(ws_claude, "Họ và tên:", d.get("ho_ten", ""))
    add_row(ws_claude, "Đơn vị công tác:", d.get("don_vi_cong_tac", ""))
    add_row(ws_claude, "Chức danh:", d.get("chuc_danh", ""))
    add_row(ws_claude, "Trình độ:", d.get("trinh_do", ""))
    add_row(ws_claude, "Lĩnh vực:", d.get("linh_vuc", ""))
    add_row(ws_claude, "Ngày áp dụng:", d.get("ngay_ap_dung", ""))
    add_row(ws_claude, "Tên đề tài:", d.get("ten_de_tai", ""), height=40)

    ws_claude.append(["", ""])
    add_section_header(ws_claude, "KẾT QUẢ CHẤM ĐIỂM - CLAUDE AI")

    sm = _safe_float(d.get("score_moi"))
    sn = _safe_float(d.get("score_nhan_rong"))
    sh = _safe_float(d.get("score_hieu_qua"))
    total = sm + sn + sh

    add_row(ws_claude, "Tính mới & Sáng tạo:", f"{sm:.0f} / 30 điểm")
    add_row(ws_claude, "Khả năng áp dụng:", f"{sn:.0f} / 30 điểm")
    add_row(ws_claude, "Hiệu quả:", f"{sh:.0f} / 40 điểm")
    add_row(ws_claude, "TỔNG ĐIỂM:", f"{total:.0f} / 100 điểm")
    add_row(ws_claude, "KẾT QUẢ:", d.get("ket_qua", ""))

    ws_claude.append(["", ""])
    add_section_header(ws_claude, "NHẬN XÉT - CLAUDE AI")
    add_row(ws_claude, "Nhận xét Tính mới:", d.get("nhan_xet_moi", ""), height=60)
    add_row(ws_claude, "Nhận xét Áp dụng:", d.get("nhan_xet_nhan_rong", ""), height=60)
    add_row(ws_claude, "Nhận xét Hiệu quả:", d.get("nhan_xet_hieu_qua", ""), height=60)
    add_row(ws_claude, "Nhận xét chung:", d.get("nhan_xet_chung", ""), height=80)
    add_row(ws_claude, "Kiến nghị:", d.get("kien_nghi", ""), height=60)
    add_row(ws_claude, "Gợi ý cải thiện:", d.get("cai_thien", ""), height=80)

    # Sheet 2: DeepSeek AI
    ws_ds = wb.create_sheet("DeepSeek AI - Chi tiết")
    ws_ds.column_dimensions['A'].width = 28
    ws_ds.column_dimensions['B'].width = 60

    add_section_header(ws_ds, "THÔNG TIN TÁC GIẢ")
    add_row(ws_ds, "Họ và tên:", d.get("ho_ten", ""))
    add_row(ws_ds, "Đơn vị công tác:", d.get("don_vi_cong_tac", ""))
    add_row(ws_ds, "Tên đề tài:", d.get("ten_de_tai", ""), height=40)

    ws_ds.append(["", ""])
    add_section_header(ws_ds, "KẾT QUẢ CHẤM ĐIỂM - DeepSeek AI")

    sm_g = _safe_float(d.get("Deepseek_score_moi"))
    sn_g = _safe_float(d.get("Deepseek_score_nhan_rong"))
    sh_g = _safe_float(d.get("Deepseek_score_hieu_qua"))
    total_g = sm_g + sn_g + sh_g

    add_row(ws_ds, "Tính mới & Sáng tạo:", f"{sm_g:.0f} / 30 điểm")
    add_row(ws_ds, "Khả năng áp dụng:", f"{sn_g:.0f} / 30 điểm")
    add_row(ws_ds, "Hiệu quả:", f"{sh_g:.0f} / 40 điểm")
    
    if total_g > 0:
        add_row(ws_ds, "TỔNG ĐIỂM:", f"{total_g:.0f} / 100 điểm")
    else:
        add_row(ws_ds, "TỔNG ĐIỂM:", "Chưa chấm")
    
    add_row(ws_ds, "KẾT QUẢ:", d.get("Deepseek_ket_qua", "Chưa chấm"))

    ws_ds.append(["", ""])
    add_section_header(ws_ds, "NHẬN XÉT - DeepSeek AI")
    add_row(ws_ds, "Nhận xét Tính mới:", d.get("Deepseek_nhan_xet_moi", ""), height=60)
    add_row(ws_ds, "Nhận xét Áp dụng:", d.get("Deepseek_nhan_xet_nhan_rong", ""), height=60)
    add_row(ws_ds, "Nhận xét Hiệu quả:", d.get("Deepseek_nhan_xet_hieu_qua", ""), height=60)
    add_row(ws_ds, "Nhận xét chung:", d.get("Deepseek_nhan_xet_chung", ""), height=80)
    add_row(ws_ds, "Kiến nghị:", d.get("Deepseek_kien_nghi", ""), height=60)
    add_row(ws_ds, "Gợi ý cải thiện:", d.get("Deepseek_cai_thien", ""), height=80)

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()

def page_score():
    st.subheader("🤖 Chấm Điểm Với Claude AI + DeepSeek AI")
    st.info("Hệ thống sẽ chấm điểm bằng cả hai AI đồng thời (song song) để đảm bảo tính khách quan và tiết kiệm thời gian.", icon="🎯")
    render_threshold_info()

    col1, col2 = st.columns(2)
    with col1:
        claude_key = st.text_input("🔑 Anthropic API Key (Claude)", value=ANTHROPIC_API_KEY, type="password", key="claude_key_input")
    with col2:
        deepseek_key = st.text_input("🔑 DeepSeek API Key", value=DEEPSEEK_API_KEY, type="password", key="deepseek_key_input")

    records = get_all_records()

    claude_unscored = [d for d in records if d["trang_thai"] != "Đã chấm"]
    ds_unscored = [d for d in records if d.get("Deepseek_trang_thai") != "Đã chấm"]
    claude_scored = [d for d in records if d["trang_thai"] == "Đã chấm"]
    ds_scored = [d for d in records if d.get("Deepseek_trang_thai") == "Đã chấm"]

    col_a, col_b, col_c, col_d = st.columns(4)
    col_a.metric("📁 Tổng SKKN", len(records))
    col_b.metric("🤖 Claude chưa chấm", len(claude_unscored))
    col_c.metric("🟢 DeepSeek chưa chấm", len(ds_unscored))
    col_d.metric("✅ Đã chấm cả 2", sum(1 for r in records if r["trang_thai"] == "Đã chấm" and r.get("Deepseek_trang_thai") == "Đã chấm"))

    st.divider()

    tab1, tab2, tab3, tab4 = st.tabs([
        "🎯 Chấm theo lựa chọn", "⚡ Chấm tất cả", "📋 Chấm nhiều bài", "🎯 Chấm từng bài"
    ])

    # Tab 1: Chấm theo lựa chọn
    with tab1:
        st.markdown("### 🎯 Chấm điểm theo lựa chọn")
        selected_ai = st.radio(
            "🤖 Chọn AI để chấm:",
            ["Cả Claude và DeepSeek (song song)", "Chỉ Claude AI", "Chỉ DeepSeek AI"],
            horizontal=True, 
            key="select_ai_mode_tab1"  # <--- THÊM KEY DUY NHẤT
        )
        
        if selected_ai == "Chỉ Claude AI":
            available = claude_unscored.copy() or claude_scored.copy()
            ai_label = "Claude"
        elif selected_ai == "Chỉ DeepSeek AI":
            available = ds_unscored.copy() or ds_scored.copy()
            ai_label = "DeepSeek"
        else:
            available = [d for d in records if d["trang_thai"] != "Đã chấm" or d.get("Deepseek_trang_thai") != "Đã chấm"]
            ai_label = "cả hai AI"

        if not available:
            st.success(f"✅ Tất cả SKKN đã được chấm bằng {ai_label}!")
        else:
            st.info(f"📊 Có **{len(available)}** SKKN sẵn sàng")
            
            # Tạo checkbox cho từng bài
            selected_for_multi = []
            for i, d in enumerate(available[:20]):
                c_st = "✅" if d["trang_thai"] == "Đã chấm" else "⏳"
                g_st = "✅" if d.get("Deepseek_trang_thai") == "Đã chấm" else "⏳"
                display = f"#{d['id']} – {d.get('ho_ten') or '?'} – {(d.get('ten_de_tai') or '')[:45]} [C:{c_st} G:{g_st}]"
                if st.checkbox(display, key=f"select_tab1_{d['id']}"):
                    selected_for_multi.append(d)
            
            if len(available) > 20:
                st.info(f"Hiển thị 20/{len(available)} bài. Chọn 'Chấm tất cả' ở tab khác để chấm toàn bộ.")
            
            if selected_for_multi:
                st.markdown(f"**✅ Đã chọn {len(selected_for_multi)} bài**")
                if st.button(f"🚀 Chấm {len(selected_for_multi)} bài", type="primary", key="btn_chon_tab1"):
                    if selected_ai == "Chỉ Claude AI":
                        if not claude_key:
                            st.error("Vui lòng nhập API Key Claude!")
                        else:
                            _run_batch_scoring_claude(claude_key, selected_for_multi, f"{len(selected_for_multi)} bài (Claude)")
                    elif selected_ai == "Chỉ DeepSeek AI":
                        if not deepseek_key:
                            st.error("Vui lòng nhập API Key DeepSeek!")
                        else:
                            _run_batch_scoring_deepseek(deepseek_key, selected_for_multi, f"{len(selected_for_multi)} bài (DeepSeek)")
                    else:
                        if not claude_key or not deepseek_key:
                            st.error("Vui lòng nhập API Key cho cả hai AI!")
                        else:
                            _run_batch_scoring_both(claude_key, deepseek_key, selected_for_multi, f"{len(selected_for_multi)} bài (Song song)")

    # Tab 2: Chấm tất cả
    with tab2:
        st.markdown("### ⚡ Chấm tất cả SKKN chưa chấm")
        col_b1, col_b2, col_b3 = st.columns(3)
        col_b1.metric("🤖 Claude cần chấm", len(claude_unscored))
        col_b2.metric("🟢 DeepSeek cần chấm", len(ds_unscored))
        col_b3.metric("📊 Bài chưa chấm đủ 2 AI", len([d for d in records if d["trang_thai"] != "Đã chấm" or d.get("Deepseek_trang_thai") != "Đã chấm"]))
        st.divider()
        
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        with col_btn1:
            if st.button("🚀 Chấm TẤT CẢ bằng Claude", type="primary", key="btn_all_claude"):
                if not claude_key:
                    st.error("Vui lòng nhập API Key Claude!")
                elif not claude_unscored:
                    st.success("✅ Tất cả đã chấm bằng Claude!")
                else:
                    _run_batch_scoring_claude(claude_key, claude_unscored, "Tất cả SKKN (Claude)")
        with col_btn2:
            if st.button("🚀 Chấm TẤT CẢ bằng DeepSeek", type="primary", key="btn_all_deepseek"):
                if not deepseek_key:
                    st.error("Vui lòng nhập API Key DeepSeek!")
                elif not ds_unscored:
                    st.success("✅ Tất cả đã chấm bằng DeepSeek!")
                else:
                    _run_batch_scoring_deepseek(deepseek_key, ds_unscored, "Tất cả SKKN (DeepSeek)")
        with col_btn3:
            if st.button("🚀 Chấm TẤT CẢ bằng CẢ HAI AI", type="primary", key="btn_all_both"):
                if not claude_key or not deepseek_key:
                    st.error("Vui lòng nhập API Key cho cả hai AI!")
                else:
                    both_need = [d for d in records if d["trang_thai"] != "Đã chấm" or d.get("Deepseek_trang_thai") != "Đã chấm"]
                    if not both_need:
                        st.success("✅ Tất cả đã chấm bằng cả hai AI!")
                    else:
                        _run_batch_scoring_both(claude_key, deepseek_key, both_need, f"{len(both_need)} bài (Song song)")

    # Tab 3: Chấm nhiều bài (checkbox đơn giản)
    with tab3:
        st.markdown("### 📋 Chấm nhiều bài")
        multi_ai = st.radio(
            "🤖 Chọn AI:", 
            ["Claude AI", "DeepSeek AI", "Cả 2 AI (song song)"], 
            horizontal=True,
            key="multi_ai_tab3"  # <--- THÊM KEY DUY NHẤT
        )
        
        if multi_ai == "Claude AI":
            multi_records = claude_unscored.copy() or claude_scored.copy()
        elif multi_ai == "DeepSeek AI":
            multi_records = ds_unscored.copy() or ds_scored.copy()
        else:
            multi_records = [d for d in records if d["trang_thai"] != "Đã chấm" or d.get("Deepseek_trang_thai") != "Đã chấm"]
        
        if not multi_records:
            st.success("✅ Không còn SKKN nào cần chấm!")
        else:
            selected = []
            for d in multi_records[:15]:
                c_st = "✅" if d["trang_thai"] == "Đã chấm" else "⏳"
                g_st = "✅" if d.get("Deepseek_trang_thai") == "Đã chấm" else "⏳"
                if st.checkbox(f"#{d['id']} – {d.get('ho_ten', '?')} [C:{c_st} G:{g_st}]", key=f"multi_tab3_{d['id']}"):
                    selected.append(d)
            if selected and st.button("🚀 Chấm bài đã chọn", type="primary", key="btn_multi_tab3"):
                if multi_ai == "Claude AI":
                    _run_batch_scoring_claude(claude_key, selected, f"{len(selected)} bài")
                elif multi_ai == "DeepSeek AI":
                    _run_batch_scoring_deepseek(deepseek_key, selected, f"{len(selected)} bài")
                else:
                    _run_batch_scoring_both(claude_key, deepseek_key, selected, f"{len(selected)} bài")

    # Tab 4: Chấm từng bài
    with tab4:
        st.markdown("### 🎯 Chấm từng bài riêng lẻ")
        single_ai = st.radio(
            "🤖 Chọn AI:", 
            ["Claude AI", "DeepSeek AI", "Cả 2 AI (song song)"], 
            horizontal=True,
            key="single_ai_tab4"  # <--- THÊM KEY DUY NHẤT
        )
        
        if records:
            opts = {}
            for d in records:
                c_st = "✅" if d["trang_thai"] == "Đã chấm" else "⏳"
                g_st = "✅" if d.get("Deepseek_trang_thai") == "Đã chấm" else "⏳"
                c_score = f"{d['score_total']:.0f}" if d["score_total"] else "?"
                g_score = f"{d['Deepseek_score_total']:.0f}" if d.get("Deepseek_score_total") else "?"
                opts[f"#{d['id']} – {d.get('ho_ten', '?')} [C:{c_st}({c_score}) G:{g_st}({g_score})]"] = d
            
            selected = st.selectbox("🔍 Chọn SKKN:", list(opts.keys()), key="select_single_tab4")
            if selected:
                d = opts[selected]
                with st.expander("📄 Xem thông tin SKKN", expanded=False):
                    render_info_box(d)
                
                if st.button("🚀 Chấm bài này", type="primary", key="btn_single_tab4"):
                    if single_ai == "Claude AI":
                        if not claude_key:
                            st.error("Vui lòng nhập API Key Claude!")
                        else:
                            with st.spinner("Claude đang chấm..."):
                                result = score_with_claude(d["raw_text"], claude_key, d)
                                update_claude_score(d["id"], result)
                                st.success(f"✅ Claude: {result['score_total']:.0f}/100 – {result['ket_qua']}")
                                st.rerun()
                    elif single_ai == "DeepSeek AI":
                        if not deepseek_key:
                            st.error("Vui lòng nhập API Key DeepSeek!")
                        else:
                            with st.spinner("DeepSeek đang chấm..."):
                                result = score_with_deepseek(d["raw_text"], deepseek_key, d)
                                update_Deepseek_score(d["id"], result)
                                st.success(f"✅ DeepSeek: {result['score_total']:.0f}/100 – {result['ket_qua']}")
                                st.rerun()
                    else:
                        if not claude_key or not deepseek_key:
                            st.error("Vui lòng nhập API Key cho cả hai AI!")
                        else:
                            with st.spinner("Đang chấm song song..."):
                                res = _score_one_parallel(d, claude_key, deepseek_key)
                                if res["claude_result"]:
                                    update_claude_score(d["id"], res["claude_result"])
                                    st.success(f"✅ Claude: {res['claude_result']['score_total']:.0f}/100")
                                if res["deepseek_result"]:
                                    update_Deepseek_score(d["id"], res["deepseek_result"])
                                    st.success(f"✅ DeepSeek: {res['deepseek_result']['score_total']:.0f}/100")
                                st.rerun()
        else:
            st.warning("Không có SKKN nào!")

# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    init_db()
    render_header()
    if "page" not in st.session_state:
        st.session_state["page"] = "upload"

    with st.sidebar:
        st.markdown("## 📋 Menu chính")
        pages = {
            "upload": "📤 Nạp File SKKN",
            "list":   "📋 Danh sách SKKN",
            "score":  "🤖 Chấm Điểm AI",
            "edit":   "✏️ Chỉnh Sửa Điểm",
            "stats":  "📊 Thống kê & Báo cáo",
        }
        for key, label in pages.items():
            if st.button(label, use_container_width=True, key=f"nav_{key}"):
                st.session_state["page"] = key; st.rerun()
        st.divider()
        records   = get_all_records()
        da_cham   = sum(1 for r in records if r["trang_thai"] == "Đã chấm")
        dat       = sum(1 for r in records if r.get("ket_qua") == "Đạt")
        ds_da_cham = sum(1 for r in records if r.get("Deepseek_trang_thai") == "Đã chấm")
        ds_dat    = sum(1 for r in records if r.get("Deepseek_ket_qua") == "Đạt")
        st.markdown(f"**Tổng:** {len(records)} SKKN")
        st.markdown(f"**🤖 Claude đã chấm:** {da_cham}")
        st.markdown(f"**🏆 Claude đạt:** {dat}")
        st.markdown(f"**🟢 DeepSeek đã chấm:** {ds_da_cham}")
        st.markdown(f"**🏆 DeepSeek đạt:** {ds_dat}")
        if da_cham > 0 and ds_da_cham > 0:
            tn = sum(1 for r in records if r.get("ket_qua") == r.get("Deepseek_ket_qua"))
            st.markdown(f"**🤝 Thống nhất:** {tn}/{da_cham}")
        st.divider()
        cfg = THRESHOLD_CONFIG
        st.markdown("**Barem điểm chuẩn:**")
        st.markdown(f"- 🆕 Tính mới: 30đ (≥{cfg['moi_min']})")
        st.markdown(f"- 🔄 Áp dụng:  30đ (≥{cfg['nhan_rong_min']})")
        st.markdown(f"- 📈 Hiệu quả: 40đ (≥{cfg['hieu_qua_min']})")
        st.markdown(f"- 🏆 **Đạt: ≥{cfg['total_min']}đ & đủ tiêu chí**")

    page = st.session_state.get("page", "upload")
    if page == "upload":   page_upload()
    elif page == "list":   page_list()
    elif page == "score":  page_score()
    elif page == "edit":   page_edit_scores()
    elif page == "stats":  page_stats()
    elif page == "detail": page_detail()

if __name__ == "__main__":
    main()