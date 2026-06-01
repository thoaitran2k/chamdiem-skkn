import streamlit as st
import pandas as pd
from docx import Document
import re
from io import BytesIO

st.set_page_config(page_title="Kiểm tra cấu trúc file SKKN", layout="wide")

st.title("🔍 Chẩn đoán lỗi cấu trúc File SKKN")
st.markdown("""
### 🔬 Công cụ debug - Tìm ra lý do không trích xuất được tên tác giả
""")

def get_text_from_doc(doc):
    """Lấy text từ docx, giữ nguyên cấu trúc"""
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                full_text.append(row_text)
    return "\n".join(full_text)

def debug_parse_author(full_text: str, lines: list) -> dict:
    """Debug chi tiết việc trích xuất tên tác giả"""
    debug_info = {
        "found_by_date_pipe": None,
        "found_by_date_no_pipe": None,
        "found_by_label": None,
        "found_by_first_lines": None,
        "all_candidates": [],
        "raw_lines_with_date": []
    }
    
    # Tìm tất cả dòng có ngày tháng
    for i, line in enumerate(lines):
        if re.search(r"\d{2}/\d{2}/\d{4}", line):
            debug_info["raw_lines_with_date"].append({
                "line_num": i,
                "content": line[:200]
            })
            
            # Trường hợp có dấu |
            if "|" in line:
                cols = [c.strip() for c in line.split("|") if c.strip()]
                debug_info["found_by_date_pipe"] = {
                    "cols": cols,
                    "possible_names": [c for c in cols if len(c) > 5 and re.search(r"[A-ZÀ-Ỹ]", c) and not re.match(r"\d", c) and not re.match(r"\d{2}/\d{2}/\d{4}", c)]
                }
            
            # Trường hợp không có dấu |
            else:
                # Thử tách bằng khoảng trắng
                parts = re.split(r'[\t\s]{2,}', line)
                if len(parts) >= 3:
                    debug_info["found_by_date_no_pipe"] = {
                        "parts": parts,
                        "possible_names": [p for p in parts if len(p) > 5 and re.search(r"[A-ZÀ-Ỹ]", p) and not re.match(r"\d{2}/\d{2}/\d{4}", p)]
                    }
    
    # Tìm dòng có "Họ và tên:"
    for i, line in enumerate(lines[:30]):
        if re.search(r"(họ\s+và\s+tên|tác\s*giả|người\s+thực\s+hiện)", line, re.IGNORECASE):
            match = re.search(r'[:;]\s*([^\n]{5,50})', line)
            if match:
                debug_info["found_by_label"] = {
                    "line": line,
                    "extracted": match.group(1).strip()
                }
            break
    
    # Tìm ở 5 dòng đầu
    for i in range(min(5, len(lines))):
        line = lines[i].strip()
        if len(line) > 10 and len(line) < 60:
            words = line.split()
            if len(words) >= 2:
                # Kiểm tra có chữ hoa
                has_upper = any(w[0].isupper() or w[0] in "ÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚÝĂĐƠƯ" for w in words if w)
                if has_upper and not any(k in line.lower() for k in ["đề tài", "sáng kiến", "báo cáo"]):
                    debug_info["found_by_first_lines"] = {
                        "line_num": i,
                        "content": line
                    }
                    debug_info["all_candidates"].append(line)
    
    return debug_info

uploaded_files = st.file_uploader("Tải các file .docx lên để debug", accept_multiple_files=True, type=['docx'])

if uploaded_files:
    for file in uploaded_files:
        st.subheader(f"📄 Debug: {file.name}")
        
        try:
            doc = Document(BytesIO(file.read()))
            full_text = get_text_from_doc(doc)
            lines = full_text.split("\n")
            
            # Hiển thị 20 dòng đầu
            with st.expander("📄 Xem 20 dòng đầu của file", expanded=True):
                for i, line in enumerate(lines[:20]):
                    st.text(f"Dòng {i+1}: {line[:150]}")
            
            # Debug trích xuất tên
            debug = debug_parse_author(full_text, lines)
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("#### 🔍 Tìm dòng có ngày tháng:")
                if debug["raw_lines_with_date"]:
                    for item in debug["raw_lines_with_date"]:
                        st.code(f"Dòng {item['line_num']+1}: {item['content']}")
                else:
                    st.warning("❌ Không tìm thấy dòng nào có ngày tháng (DD/MM/YYYY)")
            
            with col2:
                st.markdown("#### 🎯 Phát hiện tên tác giả:")
                
                if debug["found_by_date_pipe"]:
                    st.success(f"✅ Tìm thấy trong bảng có dấu |")
                    st.write("Các cột:", debug["found_by_date_pipe"]["cols"])
                    if debug["found_by_date_pipe"]["possible_names"]:
                        st.write("**Tên khả nghi:**", debug["found_by_date_pipe"]["possible_names"])
                
                elif debug["found_by_date_no_pipe"]:
                    st.info(f"⚠️ Có ngày tháng nhưng không có dấu |")
                    st.write("Các phần:", debug["found_by_date_no_pipe"]["parts"])
                
                elif debug["found_by_label"]:
                    st.success(f"✅ Tìm thấy qua label 'Họ và tên:'")
                    st.write(f"Trích xuất: {debug['found_by_label']['extracted']}")
                
                elif debug["found_by_first_lines"]:
                    st.info(f"💡 Có thể tên ở đầu file:")
                    st.write(debug["found_by_first_lines"]["content"])
                
                else:
                    st.error("❌ KHÔNG TÌM THẤY TÊN TÁC GIẢ bằng bất kỳ phương pháp nào!")
            
            # Hiển thị gợi ý sửa lỗi cụ thể
            st.markdown("#### 💡 Chẩn đoán và hướng xử lý:")
            
            if not debug["raw_lines_with_date"]:
                st.error("""
                **Nguyên nhân:** File không có dòng chứa ngày tháng (định dạng DD/MM/YYYY)
                
                **Cách sửa:** Thêm bảng thông tin tác giả với cột Ngày sinh (ví dụ: 01/01/1990)
                """)
            
            elif debug["raw_lines_with_date"] and not debug["found_by_date_pipe"]:
                st.warning("""
                **Nguyên nhân:** Có ngày tháng nhưng cấu trúc không có dấu `|` để phân tách cột
                
                **Cách sửa:** 
                1. Trong Word, tạo bảng (Table) thay vì gõ text thông thường
                2. Hoặc đảm bảo các cột cách nhau bằng dấu `|`
                """)
            
            elif debug["found_by_date_pipe"] and not debug["found_by_date_pipe"]["possible_names"]:
                st.warning("""
                **Nguyên nhân:** Có dấu `|` và ngày tháng nhưng không tìm thấy cột nào chứa tên
                
                **Cách sửa:** Đảm bảo bảng có cột "Họ và tên" với giá trị không phải số, không phải ngày tháng
                """)
            
            # Hiển thị thông tin đã trích xuất được
            st.markdown("#### 📋 Thử trích xuất bằng logic hiện tại:")
            
            # Copy logic từ app.py
            from app import parse_info_from_text
            info = parse_info_from_text(full_text, file.name)
            
            result_df = pd.DataFrame([
                {"Trường": "Họ tên", "Kết quả": info["ho_ten"] or "❌ TRỐNG"},
                {"Trường": "Đơn vị", "Kết quả": info["don_vi_cong_tac"] or "❌ TRỐNG"},
                {"Trường": "Chức danh", "Kết quả": info["chuc_danh"] or "❌ TRỐNG"},
                {"Trường": "Trình độ", "Kết quả": info["trinh_do"] or "❌ TRỐNG"},
                {"Trường": "Tên đề tài", "Kết quả": info["ten_de_tai"][:80] + "..." if len(info["ten_de_tai"]) > 80 else info["ten_de_tai"] or "❌ TRỐNG"},
            ])
            st.dataframe(result_df, hide_index=True, use_container_width=True)
            
            st.divider()
            
        except Exception as e:
            st.error(f"Lỗi đọc file: {e}")